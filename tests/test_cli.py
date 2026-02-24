from pathlib import Path
from docx import Document
from click.testing import CliRunner
from zipfile import ZipFile
from lxml import etree

from docxlate.cli import cli
from docxlate.handlers import latex


def _reset_router():
    latex.reset_document()
    latex.context.clear()


def _write_minimal_bcf(path: Path):
    path.write_text(
        """<?xml version="1.0" encoding="UTF-8"?>
<bcf:controlfile xmlns:bcf="https://sourceforge.net/projects/biblatex">
  <bcf:section number="0">
    <bcf:citekey order="1">KeyA</bcf:citekey>
    <bcf:citekey order="2">KeyB</bcf:citekey>
  </bcf:section>
  <bcf:datamodel>
    <bcf:fields>
      <bcf:field>note</bcf:field>
      <bcf:field>pubstate</bcf:field>
      <bcf:field>isbn</bcf:field>
      <bcf:field>url</bcf:field>
    </bcf:fields>
  </bcf:datamodel>
</bcf:controlfile>
""",
        encoding="utf-8",
    )


def test_cli_bridge_runs(tmp_path):
    _reset_router()
    tex_path = Path('test.tex')
    latex.context['tex_path'] = str(tex_path)
    latex.run(tex_path.read_text())
    out_path = tmp_path / 'output.docx'
    latex.save(out_path)
    assert out_path.exists()


def test_cli_uses_template_docx(tmp_path):
    runner = CliRunner()
    tex_path = tmp_path / "input.tex"
    tex_path.write_text(r"\section{Intro} Hello from input.")

    template_path = tmp_path / "template.docx"
    template_doc = Document()
    template_doc.add_paragraph("TEMPLATE MARKER")
    template_doc.save(template_path)

    out_path = tmp_path / "output.docx"
    result = runner.invoke(
        cli,
        ["convert", str(tex_path), "-o", str(out_path), "--template", str(template_path)],
    )

    assert result.exit_code == 0
    assert out_path.exists()
    rendered = Document(out_path)
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "TEMPLATE MARKER" not in text
    assert "Intro" in text


def test_cli_loads_yaml_config_for_bibliography_template(tmp_path):
    runner = CliRunner()
    tex_path = tmp_path / "doc.tex"
    tex_path.write_text(r"\cite{KeyA}.")
    (tmp_path / "doc.aux").write_text(r"\abx@aux@cite{0}{KeyA}")
    (tmp_path / "doc.bbl").write_text(Path("tests/fixtures/bbl/sample.bbl").read_text())

    config_path = tmp_path / "docxlate.yaml"
    config_path.write_text(
        "\n".join(
            [
                "bibliography_numbering: none",
                "bibliography_template: |",
                "  <% if authors %><< authors|join(\", \") >>.<% endif %>",
                "  <% if fields.journaltitle %> \\textit{<< fields.journaltitle >>}.<% endif %>",
            ]
        ),
        encoding="utf-8",
    )

    out_path = tmp_path / "output.docx"
    result = runner.invoke(
        cli,
        ["convert", str(tex_path), "-o", str(out_path), "--config", str(config_path)],
    )

    assert result.exit_code == 0
    rendered = Document(out_path)
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "A sample article" not in text
    assert "Doe, Jane, Roe, John." in text


def test_cli_rejects_invalid_yaml_config(tmp_path):
    runner = CliRunner()
    tex_path = tmp_path / "doc.tex"
    tex_path.write_text("Hello")
    config_path = tmp_path / "invalid.yaml"
    config_path.write_text("unknown_option: true\n", encoding="utf-8")

    result = runner.invoke(
        cli,
        ["convert", str(tex_path), "--config", str(config_path)],
    )

    assert result.exit_code != 0
    assert "Invalid config" in result.output


def _extract_styles_xml(docx_path: Path, out_path: Path):
    with ZipFile(docx_path, "r") as zf:
        out_path.write_bytes(zf.read("word/styles.xml"))


def _style_based_on(styles_xml_path: Path, style_id: str) -> str | None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = etree.fromstring(styles_xml_path.read_bytes())
    nodes = root.xpath(f"//w:style[@w:styleId='{style_id}']", namespaces=ns)
    if not nodes:
        return None
    value = nodes[0].xpath("string(w:basedOn/@w:val)", namespaces=ns)
    return value or None


def _set_style_based_on(styles_xml_path: Path, style_id: str, based_on: str):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = etree.fromstring(styles_xml_path.read_bytes())
    nodes = root.xpath(f"//w:style[@w:styleId='{style_id}']", namespaces=ns)
    if not nodes:
        return
    node = nodes[0]
    based_nodes = node.xpath("w:basedOn", namespaces=ns)
    if based_nodes:
        based_nodes[0].set(f"{{{ns['w']}}}val", based_on)
    else:
        based = etree.Element(f"{{{ns['w']}}}basedOn")
        based.set(f"{{{ns['w']}}}val", based_on)
        node.insert(1, based)
    styles_xml_path.write_bytes(
        etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    )


def test_cli_dump_styles_writes_styles_xml(tmp_path):
    runner = CliRunner()
    source_docx = tmp_path / "source.docx"
    Document().save(source_docx)
    dumped = tmp_path / "styles.xml"
    result = runner.invoke(
        cli,
        ["dump-styles", str(source_docx), "-o", str(dumped)],
    )
    assert result.exit_code == 0
    assert dumped.exists()
    text = dumped.read_text(encoding="utf-8")
    assert text.startswith("<?xml")
    root = etree.fromstring(text.encode("utf-8"))
    assert etree.QName(root).localname == "styles"
    assert "\n  <" in text


def test_cli_dump_theme_writes_theme_xml(tmp_path):
    runner = CliRunner()
    source_docx = tmp_path / "source.docx"
    Document().save(source_docx)
    dumped = tmp_path / "theme1.xml"
    result = runner.invoke(
        cli,
        ["dump-theme", str(source_docx), "-o", str(dumped)],
    )
    assert result.exit_code == 0
    assert dumped.exists()
    text = dumped.read_text(encoding="utf-8")
    root = etree.fromstring(text.encode("utf-8"))
    assert etree.QName(root).localname == "theme"


def test_cli_dump_font_table_writes_xml(tmp_path):
    runner = CliRunner()
    source_docx = tmp_path / "source.docx"
    Document().save(source_docx)
    dumped = tmp_path / "fontTable.xml"
    result = runner.invoke(
        cli,
        ["dump-font-table", str(source_docx), "-o", str(dumped)],
    )
    assert result.exit_code == 0
    assert dumped.exists()
    text = dumped.read_text(encoding="utf-8")
    root = etree.fromstring(text.encode("utf-8"))
    assert etree.QName(root).localname == "fonts"


def test_cli_injects_styles_xml(tmp_path):
    runner = CliRunner()
    tex_path = tmp_path / "doc.tex"
    tex_path.write_text(r"\href{https://example.com}{x}")

    base_docx = tmp_path / "base.docx"
    Document().save(base_docx)
    styles_xml = tmp_path / "styles.xml"
    _extract_styles_xml(base_docx, styles_xml)
    _set_style_based_on(styles_xml, "Caption", "Heading1")

    out_docx = tmp_path / "out.docx"
    result = runner.invoke(
        cli,
        ["convert", str(tex_path), "-o", str(out_docx), "--styles-xml", str(styles_xml)],
    )
    assert result.exit_code == 0

    extracted = tmp_path / "out_styles.xml"
    _extract_styles_xml(out_docx, extracted)
    assert _style_based_on(extracted, "Caption") == "Heading1"


def test_cli_template_accepts_styles_xml_shorthand(tmp_path):
    runner = CliRunner()
    tex_path = tmp_path / "doc.tex"
    tex_path.write_text("Hello")

    base_docx = tmp_path / "base.docx"
    Document().save(base_docx)
    styles_xml = tmp_path / "styles.xml"
    _extract_styles_xml(base_docx, styles_xml)
    _set_style_based_on(styles_xml, "Caption", "Heading1")

    out_docx = tmp_path / "out.docx"
    result = runner.invoke(
        cli,
        ["convert", str(tex_path), "-o", str(out_docx), "--template", str(styles_xml)],
    )
    assert result.exit_code == 0

    extracted = tmp_path / "out_styles.xml"
    _extract_styles_xml(out_docx, extracted)
    assert _style_based_on(extracted, "Caption") == "Heading1"


def test_cli_template_chain_applies_docx_then_xml_override(tmp_path):
    runner = CliRunner()
    tex_path = tmp_path / "doc.tex"
    tex_path.write_text("Hello")

    template_docx = tmp_path / "template.docx"
    Document().save(template_docx)
    styles_xml = tmp_path / "styles.xml"
    _extract_styles_xml(template_docx, styles_xml)
    _set_style_based_on(styles_xml, "Caption", "Heading1")

    out_docx = tmp_path / "out.docx"
    result = runner.invoke(
        cli,
        [
            "convert",
            str(tex_path),
            "-o",
            str(out_docx),
            "-t",
            str(template_docx),
            "-t",
            str(styles_xml),
        ],
    )
    assert result.exit_code == 0

    extracted = tmp_path / "out_styles.xml"
    _extract_styles_xml(out_docx, extracted)
    assert _style_based_on(extracted, "Caption") == "Heading1"


def test_cli_check_bcf_reports_present_fields(tmp_path):
    runner = CliRunner()
    bcf_path = tmp_path / "main.bcf"
    _write_minimal_bcf(bcf_path)

    result = runner.invoke(
        cli,
        ["check-bcf", str(bcf_path), "--field", "note", "--field", "isbn"],
    )
    assert result.exit_code == 0
    assert "Citekeys with order: 2" in result.output
    assert "OK: note" in result.output
    assert "OK: isbn" in result.output


def test_cli_check_bcf_fails_for_missing_field(tmp_path):
    runner = CliRunner()
    bcf_path = tmp_path / "main.bcf"
    _write_minimal_bcf(bcf_path)

    result = runner.invoke(
        cli,
        ["check-bcf", str(bcf_path), "--field", "doi", "--field", "pubstate"],
    )
    assert result.exit_code != 0
    assert "MISSING: doi" in result.output
    assert "OK: pubstate" in result.output
