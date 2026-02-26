from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from docxlate.handlers import latex


def _make_template_with_body_styles(path: Path) -> Path:
    doc = Document()
    styles = doc.styles

    try:
        body_text = styles["Body Text"]
    except KeyError:
        body_text = styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
        body_text.base_style = styles["Normal"]

    try:
        first_para = styles["First Paragraph"]
    except KeyError:
        first_para = styles.add_style("First Paragraph", WD_STYLE_TYPE.PARAGRAPH)
        first_para.base_style = body_text

    doc.add_paragraph("Template content to clear")
    doc.save(path)
    return path


def _make_template_with_bold_italic_body(path: Path) -> Path:
    doc = Document()
    styles = doc.styles
    try:
        body_text = styles["Body Text"]
    except KeyError:
        body_text = styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
        body_text.base_style = styles["Normal"]
    body_text.font.bold = True
    body_text.font.italic = True
    doc.save(path)
    return path


def test_proposal_template_applies_first_and_body_text_styles(tmp_path):
    template = _make_template_with_body_styles(tmp_path / "template_styles.docx")
    latex.reset_document(template)

    latex.run("First paragraph.\n\nSecond paragraph.")

    nonempty = [p for p in latex.doc.paragraphs if p.text.strip()]
    assert len(nonempty) >= 2
    assert nonempty[0].style.style_id == "FirstParagraph"
    assert nonempty[1].style.style_id == "BodyText"


def test_section_resets_next_body_paragraph_to_first_paragraph_style(tmp_path):
    template = _make_template_with_body_styles(tmp_path / "template_styles.docx")
    latex.reset_document(template)

    latex.run("\\section{Intro} First paragraph after heading.\n\nSecond paragraph after heading.")

    nonempty = [p for p in latex.doc.paragraphs if p.text.strip()]
    assert len(nonempty) >= 3
    assert nonempty[0].style.style_id == "Heading1"
    assert nonempty[1].style.style_id == "FirstParagraph"
    assert nonempty[2].style.style_id == "BodyText"


def test_template_bold_italic_body_style_can_be_explicitly_reset(tmp_path):
    template = _make_template_with_bold_italic_body(tmp_path / "template_bold_italic.docx")
    latex.reset_document(template)

    latex.run(r"{\mdseries\upshape reset} {\normalfont clear}")

    para = next((p for p in latex.doc.paragraphs if p.text.strip()), None)
    assert para is not None
    xml = para._element.xml
    assert "reset clear" in para.text
    assert '<w:b w:val="0"' in xml
    assert '<w:i w:val="0"' in xml
