from pathlib import Path

from docxlate.aux import parse_abx_aux_cite_order
from docxlate.handlers import latex
from docx import Document
import pytest


def _revtex_bbl_sample_text() -> str:
    return r"""
\begin{thebibliography}{2}%
\bibitem [{\citenamefont {Doe}(2024)}]{KeyA}%
  \BibitemOpen
  \bibfield {author} {\bibinfo {author} {\bibfnamefont {A.}~\bibnamefont {Doe}},\ }
  \bibfield {title} {\bibinfo {title} {Sample title A},\ }
  \href {https://doi.org/10.1000/a} {\bibfield {journal} {\bibinfo {journal} {J. A}\ }\textbf {\bibinfo {volume} {1}},\ \bibinfo {pages} {1} (\bibinfo {year} {2024})}\BibitemShut {NoStop}%
\bibitem [{\citenamefont {Roe}(2025)}]{KeyB}%
  \BibitemOpen
  \bibfield {author} {\bibinfo {author} {\bibfnamefont {B.}~\bibnamefont {Roe}},\ }
  \bibfield {title} {\bibinfo {title} {Sample title B},\ }
  \bibfield {journal} {\bibinfo {journal} {J. B}\ }\textbf {\bibinfo {volume} {2}},\ \bibinfo {pages} {11} (\bibinfo {year} {2025})\BibitemShut {NoStop}%
\end{thebibliography}
""".strip()


def test_cite_single_key_maps_to_bracket_number():
    latex.context["cite_order"] = {"Foo2025": 12}

    latex.run(r"See \cite{Foo2025}.")

    text = latex.doc.paragraphs[0].text
    assert "[12]" in text
    assert "Foo2025" not in text


def test_cite_multiple_keys_keep_order():
    latex.context["cite_order"] = {"A": 3, "B": 9}

    latex.run(r"\cite{A,B}")

    assert "[3,9]" in latex.doc.paragraphs[0].text


def test_cite_missing_key_falls_back_to_label():
    latex.context["cite_order"] = {}

    latex.run(r"\cite{UnknownKey}")

    assert "[UnknownKey]" in latex.doc.paragraphs[0].text


def test_cite_falls_back_to_bibcite_number_when_cite_order_missing():
    latex.context["cite_order"] = {}
    latex.context["bibcites"] = {"Foo2025": {"ref_num": "12"}}

    latex.run(r"See \cite{Foo2025}.")

    text = latex.doc.paragraphs[0].text
    assert "[12]" in text
    assert "Foo2025" not in text


def test_cite_text_is_present_when_mixed_with_textbf():
    latex.context["cite_order"] = {"Foo2025": 12}

    latex.run(r"\textbf{See \cite{Foo2025} now}")

    para = latex.doc.paragraphs[0]
    assert "See [12] now" in para.text
    # Ensure surrounding bold styling still appears in the paragraph.
    assert any(run.bold for run in para.runs)


def test_parse_abx_aux_cite_order_tracks_first_seen_order(tmp_path):
    aux_path = tmp_path / "sample.aux"
    aux_path.write_text(
        "\n".join(
            [
                r"\abx@aux@cite{0}{Foo}",
                r"\abx@aux@cite{0}{Bar}",
                r"\abx@aux@cite{0}{Foo}",
                r"\abx@aux@cite{0}{Baz}",
            ]
        )
    )
    order = parse_abx_aux_cite_order(aux_path)
    assert order == {"Foo": 1, "Bar": 2, "Baz": 3}


def test_cite_avoids_repeated_numbers_when_keys_are_distinct():
    # Simulates a bad upstream mapping collision.
    latex.context["cite_order"] = {"A": 3, "B": 3, "C": 3}
    latex.run(r"\cite{A,B,C}")
    assert "[3]" in latex.doc.paragraphs[0].text


def test_references_section_appended_from_bbl(tmp_path):
    tex_path = tmp_path / "doc.tex"
    aux_path = tmp_path / "doc.aux"
    bbl_path = tmp_path / "doc.bbl"

    tex_path.write_text(r"\cite{KeyA} and \cite{KeyB}.")
    aux_path.write_text(
        "\n".join(
            [
                r"\abx@aux@cite{0}{KeyA}",
                r"\abx@aux@cite{0}{KeyB}",
            ]
        )
    )
    bbl_path.write_text(Path("tests/fixtures/bbl/sample.bbl").read_text())

    latex.context["tex_path"] = str(tex_path)
    latex.run(tex_path.read_text())

    text = "\n".join(p.text for p in latex.doc.paragraphs if p.text.strip())
    assert "References" in text
    assert "A sample article" in text
    assert "Another article" in text

    xml = "\n".join(p._element.xml for p in latex.doc.paragraphs)
    assert "w:bookmarkStart" in xml
    assert "ref_bib_KeyA" in xml
    assert "ref_bib_KeyB" in xml
    assert "<w:hyperlink" in xml
    assert 'w:anchor="ref_bib_KeyA"' in xml or 'w:anchor="ref_bib_KeyB"' in xml


def test_references_missing_entry_default_policy_renders_key_row(tmp_path):
    tex_path = tmp_path / "doc.tex"
    aux_path = tmp_path / "doc.aux"
    bbl_path = tmp_path / "doc.bbl"

    tex_path.write_text(r"\cite{KeyA,MissingX,KeyB}.")
    aux_path.write_text(
        "\n".join(
            [
                r"\abx@aux@cite{0}{KeyA}",
                r"\abx@aux@cite{0}{MissingX}",
                r"\abx@aux@cite{0}{KeyB}",
            ]
        )
    )
    bbl_path.write_text(Path("tests/fixtures/bbl/sample.bbl").read_text())

    latex.context["tex_path"] = str(tex_path)
    latex.run(tex_path.read_text())

    text = "\n".join(p.text for p in latex.doc.paragraphs if p.text.strip())
    assert "[2]\tMissingX" in text
    missing_para = next((p for p in latex.doc.paragraphs if "MissingX" in p.text), None)
    assert missing_para is not None
    xml = missing_para._element.xml
    assert "<w:b/>" in xml or "<w:b " in xml
    assert 'w:val="CC0000"' in xml


def test_references_missing_entry_hole_policy_renders_blank_numbered_row(tmp_path):
    tex_path = tmp_path / "doc.tex"
    aux_path = tmp_path / "doc.aux"
    bbl_path = tmp_path / "doc.bbl"

    tex_path.write_text(r"\cite{KeyA,MissingX,KeyB}.")
    aux_path.write_text(
        "\n".join(
            [
                r"\abx@aux@cite{0}{KeyA}",
                r"\abx@aux@cite{0}{MissingX}",
                r"\abx@aux@cite{0}{KeyB}",
            ]
        )
    )
    bbl_path.write_text(Path("tests/fixtures/bbl/sample.bbl").read_text())

    latex.context["tex_path"] = str(tex_path)
    latex.context.setdefault("plugins", {}).setdefault("bibliography", {})[
        "missing_entry_policy"
    ] = "hole"
    latex.run(tex_path.read_text())

    ref_paras = [p for p in latex.doc.paragraphs if p.text.startswith("[")]
    hole_para = next((p for p in ref_paras if p.text.startswith("[2]\t") and p.text.strip() == "[2]"), None)
    assert hole_para is not None
    xml = hole_para._element.xml
    assert "<w:b/>" in xml or "<w:b " in xml
    assert 'w:val="CC0000"' in xml


def test_references_section_uses_template_bibliography_style(tmp_path):
    tex_path = tmp_path / "doc.tex"
    aux_path = tmp_path / "doc.aux"
    bbl_path = tmp_path / "doc.bbl"
    template_path = tmp_path / "template.docx"

    doc = Document()
    styles = doc.styles
    if "Bibliography" not in [s.name for s in styles]:
        styles.add_style("Bibliography", 1)
    doc.save(template_path)

    tex_path.write_text(r"\cite{KeyA}.")
    aux_path.write_text(r"\abx@aux@cite{0}{KeyA}")
    bbl_path.write_text(Path("tests/fixtures/bbl/sample.bbl").read_text())

    latex.reset_document(str(template_path))
    latex.context["tex_path"] = str(tex_path)
    latex.run(tex_path.read_text())

    refs = [p for p in latex.doc.paragraphs if "A sample article" in p.text]
    assert refs, "Expected bibliography paragraph not found"
    assert refs[0].style.name == "Bibliography"


def test_references_section_defaults_to_bracket_number_and_tab_layout(tmp_path):
    tex_path = tmp_path / "doc.tex"
    aux_path = tmp_path / "doc.aux"
    bbl_path = tmp_path / "doc.bbl"

    tex_path.write_text(r"\cite{KeyA}.")
    aux_path.write_text(r"\abx@aux@cite{0}{KeyA}")
    bbl_path.write_text(Path("tests/fixtures/bbl/sample.bbl").read_text())

    latex.context["tex_path"] = str(tex_path)
    latex.run(tex_path.read_text())

    ref_para = next((p for p in latex.doc.paragraphs if "A sample article" in p.text), None)
    assert ref_para is not None
    assert ref_para.text.startswith("[1]\t")
    xml = ref_para._element.xml
    assert "<w:tab/>" in xml
    assert "<w:tabs>" in xml
    assert 'w:hanging="' in xml


def test_references_section_can_disable_numbering_layout(tmp_path):
    tex_path = tmp_path / "doc.tex"
    aux_path = tmp_path / "doc.aux"
    bbl_path = tmp_path / "doc.bbl"

    tex_path.write_text(r"\cite{KeyA}.")
    aux_path.write_text(r"\abx@aux@cite{0}{KeyA}")
    bbl_path.write_text(Path("tests/fixtures/bbl/sample.bbl").read_text())

    latex.context["tex_path"] = str(tex_path)
    latex.context.setdefault("plugins", {}).setdefault("bibliography", {})["numbering"] = "none"
    latex.run(tex_path.read_text())

    ref_para = next((p for p in latex.doc.paragraphs if "A sample article" in p.text), None)
    assert ref_para is not None
    assert not ref_para.text.startswith("[1]\t")
    xml = ref_para._element.xml
    assert "<w:tab/>" not in xml


def test_cite_compresses_numeric_ranges():
    latex.context["cite_order"] = {
        "A": 1,
        "B": 2,
        "C": 3,
        "D": 4,
        "E": 5,
    }
    latex.context.setdefault("plugins", {}).setdefault("bibliography", {})[
        "citation_compress_ranges"
    ] = True

    latex.run(r"\cite{A,B,C,D,E}")

    assert "[1–5]" in latex.doc.paragraphs[0].text


def test_cite_compresses_with_min_run_threshold():
    latex.context["cite_order"] = {
        "A": 1,
        "B": 2,
        "C": 4,
    }
    latex.context.setdefault("plugins", {}).setdefault("bibliography", {})[
        "citation_compress_ranges"
    ] = True
    latex.context.setdefault("plugins", {}).setdefault("bibliography", {})[
        "citation_range_min_run"
    ] = 3

    latex.run(r"\cite{A,B,C}")

    # With min run 3, [1,2] should not collapse.
    assert "[1,2,4]" in latex.doc.paragraphs[0].text


@pytest.mark.xfail(reason="REVTeX note-like bibitems are not classified separately yet")
def test_cite_mixed_reference_and_note_keys_revtex_style():
    # Numeric references mixed with note-like keys should avoid folding across types.
    latex.context["cite_order"] = {
        "RefA": 1,
        "NoteX": 2,
        "RefB": 3,
    }
    latex.context.setdefault("plugins", {}).setdefault("bibliography", {})[
        "citation_compress_ranges"
    ] = True

    latex.run(r"\cite{RefA,NoteX,RefB}")

    # Desired future behavior: preserve note key separately (not folded into 1-3).
    assert "[1,NoteX,3]" in latex.doc.paragraphs[0].text


@pytest.mark.xfail(reason="Author-year citation rendering mode is not implemented yet")
def test_cite_authoryear_mode_renders_names_and_year():
    latex.context["citation_mode"] = "authoryear"
    latex.context["cite_order"] = {"Smith2020": 1}
    latex.context["refs"] = {"Smith2020": {"ref_num": "Smith, 2020"}}

    latex.run(r"\cite{Smith2020}")

    assert "(Smith, 2020)" in latex.doc.paragraphs[0].text


def test_references_section_revtex_bbl_and_bibcite_aux(tmp_path):
    tex_path = tmp_path / "doc.tex"
    aux_path = tmp_path / "doc.aux"
    bbl_path = tmp_path / "doc.bbl"

    tex_path.write_text(r"\cite{KeyA} and \cite{KeyB}.")
    aux_path.write_text(
        "\n".join(
            [
                r"\bibcite{KeyA}{{1}{2024}{{Doe}}{{Doe}}}",
                r"\bibcite{KeyB}{{2}{2025}{{Roe}}{{Roe}}}",
            ]
        )
    )
    bbl_path.write_text(_revtex_bbl_sample_text())

    latex.context["tex_path"] = str(tex_path)
    latex.run(tex_path.read_text())

    text = "\n".join(p.text for p in latex.doc.paragraphs if p.text.strip())
    assert "[1]" in text
    assert "[2]" in text
    assert "References" in text
    assert "Sample title A" in text
    assert "Sample title B" in text
