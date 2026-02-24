import base64

from docx import Document
from docx.oxml.ns import qn

from docxlate.docx_ext.floating import (
    insert_wrapped_figure_caption_group_anchor,
    convert_inline_drawing_to_wrapped_anchor,
    insert_wrapped_caption_anchor,
    next_anchor_group_id,
)


def _write_png(path):
    data = (
        b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+X2w0AAAAASUVORK5CYII="
    )
    path.write_bytes(base64.b64decode(data))


def test_insert_wrapped_caption_anchor_tags_group_membership():
    doc = Document()
    source = doc.add_paragraph("Caption text")
    host = doc.add_paragraph("Host")

    insert_wrapped_caption_anchor(
        doc,
        source_paragraph=source,
        anchor_paragraph=host,
        place="r",
        pos_y_emu=0,
        box_cx_emu=1200000,
        box_cy_emu=320000,
        group_id=7,
    )

    assert "docxlate-wrap-group:7:caption" in host._element.xml
    assert next_anchor_group_id(doc) == 8


def test_convert_inline_anchor_tags_group_membership(tmp_path):
    image_path = tmp_path / "sample.png"
    _write_png(image_path)

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(str(image_path))

    drawing = run._r.find(qn("w:drawing"))
    assert drawing is not None
    anchor = convert_inline_drawing_to_wrapped_anchor(
        drawing,
        place="r",
        pos_y_emu=0,
        group_id=12,
    )
    assert anchor is not None
    assert "docxlate-wrap-group:12:image" in para._element.xml
    assert next_anchor_group_id(doc) == 13


def test_group_anchor_image_paragraph_sets_at_least_line_height(tmp_path):
    image_path = tmp_path / "sample.png"
    _write_png(image_path)

    doc = Document()
    host = doc.add_paragraph("Host")
    image_para = doc.add_paragraph()
    image_run = image_para.add_run()
    image_run.add_picture(str(image_path))
    caption_para = doc.add_paragraph("Caption")

    insert_wrapped_figure_caption_group_anchor(
        doc,
        image_run=image_run,
        caption_paragraph=caption_para,
        anchor_paragraph=host,
        place="r",
        pos_y_emu=0,
        box_cx_emu=1200000,
        box_cy_emu=1200000,
        gap_emu=114300,
    )

    xml = host._element.xml
    assert 'w:lineRule="atLeast"' in xml
    assert 'w:line="' in xml
