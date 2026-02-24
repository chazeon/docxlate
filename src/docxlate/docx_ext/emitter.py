from __future__ import annotations
from contextlib import contextmanager
from collections.abc import Mapping

from docx.oxml.ns import qn
from docx.shared import Inches
from docxlate.model import EquationSpec, LinkTarget, TextSpan
from docxlate.utils import apply_theme_font, inject_omml
from .floating import (
    convert_inline_drawing_to_wrapped_anchor,
    insert_wrapped_figure_caption_group_anchor,
    next_anchor_group_id,
    insert_wrapped_caption_anchor,
)
from .hyperlink import HyperlinkWriter
from .run_style import apply_text_span_style


class DocxEmitterBackend:
    """
    DOCX emission backend for resolved inline spans and equation specs.
    Keeps OOXML-specific logic out of core traversal/dispatch.
    """

    def __init__(self, context: dict):
        self.context = context
        self._active_link: LinkTarget | None = None
        self._hyperlinks = HyperlinkWriter()

    def begin_paragraph(
        self,
        doc,
        *,
        role: str | None = None,
        style_name: str | None = None,
        style_table: dict | None = None,
    ):
        resolved = style_name or self._resolve_paragraph_style(doc, role, style_table)
        if resolved:
            return doc.add_paragraph(style=resolved)
        return doc.add_paragraph()

    def begin_link(self, target: LinkTarget):
        if self._active_link is not None:
            raise RuntimeError("Nested hyperlinks are not supported")
        self._active_link = target

    def end_link(self):
        self._active_link = None

    @contextmanager
    def link_scope(self, target: LinkTarget):
        self.begin_link(target)
        try:
            yield
        finally:
            self.end_link()

    def has_active_link(self) -> bool:
        return self._active_link is not None

    def emit_span(self, paragraph, span: TextSpan):
        if self._active_link is None:
            self._emit_plain_span(paragraph, span)
            return
        self._emit_linked_span(paragraph, span, self._active_link)

    def emit_line_break(self, paragraph):
        paragraph.add_run().add_break()

    def emit_equation(self, paragraph, spec: EquationSpec):
        xsl_path = self.context.get("mathml2omml_xsl_path")
        ok = inject_omml(
            paragraph,
            spec.latex,
            xsl_path=xsl_path,
            color=spec.color,
            display=spec.display,
            style=spec.style,
        )
        if not ok and not xsl_path:
            self._warn_missing_math_xsl()
        if spec.number:
            run = paragraph.add_run(f" ({spec.number})")
            apply_theme_font(run, "minor")

    def emit_image(self, paragraph, image_path: str, *, width_emu: int | None = None):
        run = paragraph.add_run()
        if width_emu is not None:
            run.add_picture(str(image_path), width=width_emu)
        else:
            run.add_picture(str(image_path))
        return run

    def convert_image_run_to_wrap_anchor(
        self,
        run,
        *,
        place: str | None,
        pos_y_emu: int = 0,
        group_id: int | None = None,
    ):
        drawing = run._r.find(qn("w:drawing"))
        if drawing is None:
            return None
        return convert_inline_drawing_to_wrapped_anchor(
            drawing,
            place=place,
            pos_y_emu=pos_y_emu,
            wrap_distances_emu=self._wrap_distances_emu(),
            group_id=group_id,
        )

    def emit_wrapped_caption_anchor(
        self,
        doc,
        *,
        source_paragraph,
        anchor_paragraph=None,
        place: str | None,
        pos_y_emu: int,
        box_cx_emu: int,
        box_cy_emu: int,
        group_id: int | None = None,
    ):
        return insert_wrapped_caption_anchor(
            doc,
            source_paragraph=source_paragraph,
            anchor_paragraph=anchor_paragraph,
            place=place,
            pos_y_emu=pos_y_emu,
            box_cx_emu=box_cx_emu,
            box_cy_emu=box_cy_emu,
            wrap_distances_emu=self._wrap_distances_emu(),
            textbox_insets_emu=self._textbox_insets_emu(),
            group_id=group_id,
        )

    def emit_wrapped_figure_caption_group_anchor(
        self,
        doc,
        *,
        image_run,
        caption_paragraph,
        anchor_paragraph=None,
        place: str | None,
        pos_y_emu: int,
        box_cx_emu: int,
        box_cy_emu: int,
        gap_emu: int,
    ):
        return insert_wrapped_figure_caption_group_anchor(
            doc,
            image_run=image_run,
            caption_paragraph=caption_paragraph,
            anchor_paragraph=anchor_paragraph,
            place=place,
            pos_y_emu=pos_y_emu,
            box_cx_emu=box_cx_emu,
            box_cy_emu=box_cy_emu,
            gap_emu=gap_emu,
            wrap_distances_emu=self._wrap_distances_emu(),
            textbox_insets_emu=self._textbox_insets_emu(),
        )

    def reserve_wrap_group_id(self, doc) -> int:
        return next_anchor_group_id(doc)

    def _wrap_distances_emu(self) -> dict[str, int]:
        def _to_emu(value, default_emu: int) -> int:
            if value is None:
                return default_emu
            try:
                inches = float(value)
            except (TypeError, ValueError):
                return default_emu
            if inches < 0:
                return default_emu
            return int(Inches(inches))

        image_cfg = None
        plugins = self.context.get("plugins")
        if isinstance(plugins, Mapping):
            figure_cfg = plugins.get("figure")
            if isinstance(figure_cfg, Mapping):
                candidate = figure_cfg.get("image")
                if isinstance(candidate, Mapping):
                    image_cfg = candidate
        wrap_cfg = None
        if isinstance(image_cfg, Mapping):
            candidate = image_cfg.get("wrap")
            if isinstance(candidate, Mapping):
                wrap_cfg = candidate

        def _side_emu(
            side_key: str,
            *,
            default_emu: int,
        ) -> int:
            box = None
            if isinstance(wrap_cfg, Mapping):
                candidate = wrap_cfg.get("pad")
                if isinstance(candidate, Mapping):
                    box = candidate
            if isinstance(box, Mapping) and side_key in box:
                return _to_emu(box.get(side_key), default_emu)
            return default_emu

        return {
            "dist_t": _side_emu(
                "top",
                default_emu=0,
            ),
            "dist_b": _side_emu(
                "bottom",
                default_emu=0,
            ),
            "dist_l": _side_emu(
                "left",
                default_emu=114300,
            ),
            "dist_r": _side_emu(
                "right",
                default_emu=114300,
            ),
        }

    def _textbox_insets_emu(self) -> dict[str, int]:
        def _to_emu(value, default_emu: int) -> int:
            if value is None:
                return default_emu
            try:
                inches = float(value)
            except (TypeError, ValueError):
                return default_emu
            if inches < 0:
                return default_emu
            return int(Inches(inches))

        image_cfg = None
        plugins = self.context.get("plugins")
        if isinstance(plugins, Mapping):
            figure_cfg = plugins.get("figure")
            if isinstance(figure_cfg, Mapping):
                candidate = figure_cfg.get("image")
                if isinstance(candidate, Mapping):
                    image_cfg = candidate
        wrap_cfg = None
        if isinstance(image_cfg, Mapping):
            candidate = image_cfg.get("wrap")
            if isinstance(candidate, Mapping):
                wrap_cfg = candidate

        def _side_emu(
            side_key: str,
            *,
            default_emu: int,
        ) -> int:
            box = None
            if isinstance(wrap_cfg, Mapping):
                candidate = wrap_cfg.get("inset")
                if isinstance(candidate, Mapping):
                    box = candidate
            if isinstance(box, Mapping) and side_key in box:
                return _to_emu(box.get(side_key), default_emu)
            return default_emu

        # Keep conservative defaults; explicit config overrides.
        return {
            "l_ins": _side_emu(
                "left",
                default_emu=0,
            ),
            "r_ins": _side_emu(
                "right",
                default_emu=0,
            ),
            "t_ins": _side_emu(
                "top",
                default_emu=0,
            ),
            "b_ins": _side_emu(
                "bottom",
                default_emu=0,
            ),
        }

    def _emit_plain_span(self, paragraph, span: TextSpan):
        run = paragraph.add_run(span.text)
        if span.char_role:
            try:
                run.style = span.char_role
            except Exception:
                pass
        apply_text_span_style(run, span)

    def _emit_linked_span(self, paragraph, span: TextSpan, link: LinkTarget):
        if not self._hyperlinks.emit_span(paragraph, span, link):
            self._emit_plain_span(paragraph, span)

    def _warn_missing_math_xsl(self):
        msg = (
            "Math OMML stylesheet path is not configured "
            "(set mathml2omml_xsl_path in config)."
        )
        warnings = self.context.setdefault("warnings", [])
        if msg not in warnings:
            warnings.append(msg)

    def _resolve_paragraph_style(self, doc, role: str | None, style_table: dict | None):
        if not role or not style_table:
            return None
        candidates = style_table.get(role, [])
        if isinstance(candidates, str):
            candidates = [candidates]
        paragraph_styles = [style for style in doc.styles if style.type == 1]
        styles_by_name = {style.name: style for style in paragraph_styles}
        style_name_by_id = {style.style_id: style.name for style in paragraph_styles}
        for candidate in candidates:
            if candidate in styles_by_name:
                return candidate
            resolved_name = style_name_by_id.get(candidate)
            if resolved_name:
                return resolved_name
        return None
