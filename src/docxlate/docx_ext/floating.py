from __future__ import annotations

from copy import deepcopy
from math import ceil

from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn

nsmap.setdefault(
    "wps", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
)


def _next_docpr_id(doc) -> int:
    ids = []
    for paragraph in doc.paragraphs:
        for drawing in paragraph._p.findall(".//" + qn("w:drawing")):
            for doc_pr in drawing.findall(".//" + qn("wp:docPr")):
                value = doc_pr.get("id")
                if value and value.isdigit():
                    ids.append(int(value))
    return max(ids, default=0) + 1


def next_anchor_group_id(doc) -> int:
    """
    Allocate a stable numeric group id used to associate related wrapped anchors
    (e.g., image + caption) before migrating to true grouped shapes.
    """
    max_group = 0
    for paragraph in doc.paragraphs:
        for drawing in paragraph._p.findall(".//" + qn("w:drawing")):
            for doc_pr in drawing.findall(".//" + qn("wp:docPr")):
                descr = doc_pr.get("descr", "")
                if not descr.startswith("docxlate-wrap-group:"):
                    continue
                payload = descr.removeprefix("docxlate-wrap-group:")
                gid, _, _role = payload.partition(":")
                if gid.isdigit():
                    max_group = max(max_group, int(gid))
    return max_group + 1


def _tag_group_member(anchor, *, group_id: int, role: str):
    doc_pr = anchor.find(qn("wp:docPr"))
    if doc_pr is None:
        return
    doc_pr.set("descr", f"docxlate-wrap-group:{int(group_id)}:{role}")
    name = doc_pr.get("name") or "Wrapped Object"
    doc_pr.set("name", f"{name} (Group {int(group_id)} {role})")


def _build_common_anchor(
    *,
    place: str | None,
    pos_y_emu: int,
    cx_emu: int,
    cy_emu: int,
    wrap_distances_emu: dict[str, int] | None = None,
):
    align = "right" if "r" in (place or "").lower() else "left"
    distances = wrap_distances_emu or {
        "dist_t": 0,
        "dist_b": 0,
        "dist_l": 114300,
        "dist_r": 114300,
    }

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", str(max(0, int(distances.get("dist_t", 0)))))
    anchor.set("distB", str(max(0, int(distances.get("dist_b", 0)))))
    anchor.set("distL", str(max(0, int(distances.get("dist_l", 114300)))))
    anchor.set("distR", str(max(0, int(distances.get("dist_r", 114300)))))
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251661312")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("hidden", "0")
    anchor.set("allowOverlap", "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "margin")
    align_elm = OxmlElement("wp:align")
    align_elm.text = align
    position_h.append(align_elm)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    pos_offset = OxmlElement("wp:posOffset")
    pos_offset.text = str(max(0, pos_y_emu))
    position_v.append(pos_offset)
    anchor.append(position_v)

    extent = OxmlElement("wp:extent")
    extent.set("cx", str(max(1, cx_emu)))
    extent.set("cy", str(max(1, cy_emu)))
    anchor.append(extent)

    effect_extent = OxmlElement("wp:effectExtent")
    effect_extent.set("l", "0")
    effect_extent.set("t", "0")
    effect_extent.set("r", "0")
    effect_extent.set("b", "0")
    anchor.append(effect_extent)

    wrap_square = OxmlElement("wp:wrapSquare")
    wrap_square.set("wrapText", "bothSides")
    wrap_square.set("distT", str(max(0, int(distances.get("dist_t", 0)))))
    wrap_square.set("distB", str(max(0, int(distances.get("dist_b", 0)))))
    wrap_square.set("distL", str(max(0, int(distances.get("dist_l", 114300)))))
    wrap_square.set("distR", str(max(0, int(distances.get("dist_r", 114300)))))
    anchor.append(wrap_square)
    return anchor


def convert_inline_drawing_to_wrapped_anchor(
    drawing_elm,
    *,
    place: str | None,
    pos_y_emu: int = 0,
    wrap_distances_emu: dict[str, int] | None = None,
    group_id: int | None = None,
):
    inline = drawing_elm.find(qn("wp:inline"))
    if inline is None:
        return None

    extent = inline.find(qn("wp:extent"))
    cx = int(extent.get("cx", "1")) if extent is not None else 1
    cy = int(extent.get("cy", "1")) if extent is not None else 1
    anchor = _build_common_anchor(
        place=place,
        pos_y_emu=pos_y_emu,
        cx_emu=cx,
        cy_emu=cy,
        wrap_distances_emu=wrap_distances_emu,
    )

    doc_pr = inline.find(qn("wp:docPr"))
    if doc_pr is not None:
        anchor.append(deepcopy(doc_pr))
    c_nv_graphic_frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))
    if c_nv_graphic_frame_pr is not None:
        anchor.append(deepcopy(c_nv_graphic_frame_pr))
    graphic = inline.find(qn("a:graphic"))
    if graphic is not None:
        anchor.append(deepcopy(graphic))

    drawing_elm.remove(inline)
    drawing_elm.append(anchor)
    if group_id is not None:
        _tag_group_member(anchor, group_id=group_id, role="image")
    return anchor


def insert_wrapped_caption_anchor(
    doc,
    *,
    source_paragraph,
    anchor_paragraph=None,
    place: str | None,
    pos_y_emu: int,
    box_cx_emu: int,
    box_cy_emu: int,
    wrap_distances_emu: dict[str, int] | None = None,
    textbox_insets_emu: dict[str, int] | None = None,
    group_id: int | None = None,
):
    anchor = _build_common_anchor(
        place=place,
        pos_y_emu=pos_y_emu,
        cx_emu=box_cx_emu,
        cy_emu=box_cy_emu,
        wrap_distances_emu=wrap_distances_emu,
    )

    doc_pr = OxmlElement("wp:docPr")
    doc_pr.set("id", str(_next_docpr_id(doc)))
    doc_pr.set("name", "Wrapped Caption")
    anchor.append(doc_pr)
    anchor.append(OxmlElement("wp:cNvGraphicFramePr"))

    graphic = OxmlElement("a:graphic")
    graphic_data = OxmlElement("a:graphicData")
    graphic_data.set(
        "uri", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    )
    graphic.append(graphic_data)
    anchor.append(graphic)

    wsp = OxmlElement("wps:wsp")
    graphic_data.append(wsp)
    wsp.append(OxmlElement("wps:cNvSpPr"))

    sp_pr = OxmlElement("wps:spPr")
    wsp.append(sp_pr)
    xfrm = OxmlElement("a:xfrm")
    sp_pr.append(xfrm)
    off = OxmlElement("a:off")
    off.set("x", "0")
    off.set("y", "0")
    xfrm.append(off)
    ext = OxmlElement("a:ext")
    ext.set("cx", str(max(1, box_cx_emu)))
    ext.set("cy", str(max(1, box_cy_emu)))
    xfrm.append(ext)
    geom = OxmlElement("a:prstGeom")
    geom.set("prst", "rect")
    geom.append(OxmlElement("a:avLst"))
    sp_pr.append(geom)
    solid_fill = OxmlElement("a:solidFill")
    srgb = OxmlElement("a:srgbClr")
    srgb.set("val", "FFFFFF")
    solid_fill.append(srgb)
    sp_pr.append(solid_fill)
    ln = OxmlElement("a:ln")
    ln.append(OxmlElement("a:noFill"))
    sp_pr.append(ln)

    txbx = OxmlElement("wps:txbx")
    wsp.append(txbx)
    txbx_content = OxmlElement("w:txbxContent")
    txbx.append(txbx_content)
    txbx_content.append(deepcopy(source_paragraph._p))
    body_pr = OxmlElement("wps:bodyPr")
    insets = textbox_insets_emu or {}
    for key, attr in (
        ("l_ins", "lIns"),
        ("r_ins", "rIns"),
        ("t_ins", "tIns"),
        ("b_ins", "bIns"),
    ):
        if key not in insets:
            continue
        body_pr.set(attr, str(max(0, int(insets[key]))))
    # Keep caption width fixed and let Word grow the box vertically as needed.
    body_pr.append(OxmlElement("a:spAutoFit"))
    wsp.append(body_pr)

    anchor_para = anchor_paragraph
    if anchor_para is None:
        anchor_para = doc.add_paragraph()
        try:
            anchor_para.style = source_paragraph.style
        except Exception:
            pass
        anchor_para.paragraph_format.space_before = 0
        anchor_para.paragraph_format.space_after = 0
    run = anchor_para.add_run()
    drawing = OxmlElement("w:drawing")
    drawing.append(anchor)
    run._r.append(drawing)

    body = doc._element.body
    body.remove(source_paragraph._p)
    if group_id is not None:
        _tag_group_member(anchor, group_id=group_id, role="caption")
    return anchor_para


def insert_wrapped_figure_caption_group_anchor(
    doc,
    *,
    image_run,
    caption_paragraph,
    anchor_paragraph=None,
    place: str | None,
    pos_y_emu: int,
    box_cx_emu: int,
    box_cy_emu: int,
    gap_emu: int,
    wrap_distances_emu: dict[str, int] | None = None,
    textbox_insets_emu: dict[str, int] | None = None,
):
    drawing_elm = image_run._r.find(qn("w:drawing"))
    if drawing_elm is None:
        return None
    inline = drawing_elm.find(qn("wp:inline"))
    if inline is None:
        return None

    anchor = _build_common_anchor(
        place=place,
        pos_y_emu=pos_y_emu,
        cx_emu=box_cx_emu,
        cy_emu=box_cy_emu,
        wrap_distances_emu=wrap_distances_emu,
    )

    doc_pr = OxmlElement("wp:docPr")
    doc_pr.set("id", str(_next_docpr_id(doc)))
    doc_pr.set("name", "Wrapped Figure+Caption")
    anchor.append(doc_pr)
    anchor.append(OxmlElement("wp:cNvGraphicFramePr"))

    graphic = OxmlElement("a:graphic")
    graphic_data = OxmlElement("a:graphicData")
    graphic_data.set(
        "uri", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    )
    graphic.append(graphic_data)
    anchor.append(graphic)

    wsp = OxmlElement("wps:wsp")
    graphic_data.append(wsp)
    wsp.append(OxmlElement("wps:cNvSpPr"))

    sp_pr = OxmlElement("wps:spPr")
    wsp.append(sp_pr)
    xfrm = OxmlElement("a:xfrm")
    sp_pr.append(xfrm)
    off = OxmlElement("a:off")
    off.set("x", "0")
    off.set("y", "0")
    xfrm.append(off)
    ext = OxmlElement("a:ext")
    ext.set("cx", str(max(1, box_cx_emu)))
    ext.set("cy", str(max(1, box_cy_emu)))
    xfrm.append(ext)
    geom = OxmlElement("a:prstGeom")
    geom.set("prst", "rect")
    geom.append(OxmlElement("a:avLst"))
    sp_pr.append(geom)
    solid_fill = OxmlElement("a:solidFill")
    srgb = OxmlElement("a:srgbClr")
    srgb.set("val", "FFFFFF")
    solid_fill.append(srgb)
    sp_pr.append(solid_fill)
    ln = OxmlElement("a:ln")
    ln.append(OxmlElement("a:noFill"))
    sp_pr.append(ln)

    txbx = OxmlElement("wps:txbx")
    wsp.append(txbx)
    txbx_content = OxmlElement("w:txbxContent")
    txbx.append(txbx_content)

    image_para = OxmlElement("w:p")
    image_run_copy = deepcopy(image_run._r)
    image_para.append(image_run_copy)
    image_para_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(max(0, int(gap_emu)) // 635))
    inline_copy = image_run_copy.find(qn("w:drawing") + "/" + qn("wp:inline"))
    if inline_copy is not None:
        image_extent = inline_copy.find(qn("wp:extent"))
        if image_extent is not None:
            cy = image_extent.get("cy", "0")
            if cy.isdigit():
                # Prevent style-driven exact line spacing from clipping tall inline images.
                line_twips = max(1, int(ceil(int(cy) / 635.0)) + 20)
                spacing.set(qn("w:line"), str(line_twips))
                spacing.set(qn("w:lineRule"), "atLeast")
    image_para_pr.append(spacing)
    image_para.insert(0, image_para_pr)
    txbx_content.append(image_para)
    txbx_content.append(deepcopy(caption_paragraph._p))

    body_pr = OxmlElement("wps:bodyPr")
    # Grouped mode currently favors clip-safety over custom insets.
    body_pr.set("lIns", "0")
    body_pr.set("rIns", "0")
    body_pr.set("tIns", "0")
    body_pr.set("bIns", "0")
    body_pr.set("vertOverflow", "overflow")
    body_pr.append(OxmlElement("a:spAutoFit"))
    wsp.append(body_pr)

    anchor_para = anchor_paragraph
    if anchor_para is None:
        anchor_para = doc.add_paragraph()
        try:
            anchor_para.style = caption_paragraph.style
        except Exception:
            pass
        anchor_para.paragraph_format.space_before = 0
        anchor_para.paragraph_format.space_after = 0
    run = anchor_para.add_run()
    drawing = OxmlElement("w:drawing")
    drawing.append(anchor)
    run._r.append(drawing)

    image_parent = image_run._r.getparent()
    if image_parent is not None:
        image_parent.remove(image_run._r)

    body = doc._element.body
    if caption_paragraph._p.getparent() is body:
        body.remove(caption_paragraph._p)
    return anchor_para
