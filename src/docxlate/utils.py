from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import RGBColor
from lxml import etree
import latex2mathml.converter
from pathlib import Path

MATHML_NAMESPACE = "http://www.w3.org/1998/Math/MathML"
OMML_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_XSLT_CACHE: dict[str, etree.XSLT] = {}
_OMML_NS = {"m": OMML_NAMESPACE}


def _get_mathml_to_omml_transform(xsl_path: str | Path):
    resolved = str(Path(xsl_path).expanduser().resolve())
    cached = _XSLT_CACHE.get(resolved)
    if cached is not None:
        return cached
    transform = etree.XSLT(etree.parse(resolved))
    _XSLT_CACHE[resolved] = transform
    return transform


def apply_theme_font(run, theme='minor'):
    """Applies Major/Minor OpenType theme attributes."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    val = 'major' if theme == 'major' else 'minor'
    rFonts.set(qn('w:asciiTheme'), f'{val}Ascii')
    rFonts.set(qn('w:hAnsiTheme'), f'{val}HAnsi')


def inject_omml(
    paragraph,
    latex_str,
    *,
    xsl_path: str | Path | None = None,
    color: str | None = None,
    display: bool = False,
):
    """Bridge LaTeX math to Word Math (OMML) via the provided stylesheet."""
    try:
        mathml = latex2mathml.converter.convert(latex_str)
        mathml_element = etree.fromstring(mathml.encode("utf-8"))
        if not xsl_path:
            run = paragraph.add_run(mathml)
            _apply_text_run_color(run, color)
            return False
        transform = _get_mathml_to_omml_transform(xsl_path)
        omml_result = transform(mathml_element)
        # Never let post-processing break core math rendering.
        try:
            _normalize_omml_script_bases(omml_result)
            _apply_native_math_run_properties(omml_result, color=color)
            _apply_native_nary_control_properties(omml_result, color=color)
        except Exception:
            pass
        omml_node = _coerce_omml_node(omml_result)
        if display:
            omml_node = _wrap_display_omath(omml_node)
        omml_xml = etree.tostring(omml_node, encoding="utf-8")
        paragraph._element.append(parse_xml(omml_xml))
        return True
    except Exception:
        # Fallback to MathML text for inspection instead of opaque error tags.
        try:
            run = paragraph.add_run(latex2mathml.converter.convert(latex_str))
            _apply_text_run_color(run, color)
        except Exception:
            run = paragraph.add_run(latex_str)
            _apply_text_run_color(run, color)
        return False


def _normalize_omml_script_bases(omml_root) -> None:
    """
    Word can render empty script bases (<m:t/>) with visible spacing holes.
    Normalize them to a zero-width placeholder to match stable producer output.
    """
    root = omml_root.getroot() if hasattr(omml_root, "getroot") else omml_root
    for text_node in root.xpath(
        ".//m:sSub/m:e//m:t | .//m:sSup/m:e//m:t | .//m:sSubSup/m:e//m:t",
        namespaces=_OMML_NS,
    ):
        if text_node.text in (None, ""):
            text_node.text = "\u200b"


def _apply_text_run_color(run, color: str | None) -> None:
    if not color:
        return
    try:
        run.font.color.rgb = RGBColor.from_string(color)
    except Exception:
        return


def _apply_native_math_run_properties(omml_root, *, color: str | None) -> None:
    """
    Match Word-native pattern observed in edited DOCX:
    each m:r carries a direct w:rPr child (with Cambria Math fonts,
    and optional w:color for scoped math coloring).
    """
    root = omml_root.getroot() if hasattr(omml_root, "getroot") else omml_root
    for m_run in root.xpath(".//m:r", namespaces=_OMML_NS):
        w_rpr = m_run.find(_wtag("rPr"))
        if w_rpr is None:
            w_rpr = etree.Element(qn("w:rPr"))
            m_run.insert(0, w_rpr)
        r_fonts = w_rpr.find(_wtag("rFonts"))
        if r_fonts is None:
            r_fonts = etree.Element(qn("w:rFonts"))
            w_rpr.append(r_fonts)
        if r_fonts.get(qn("w:ascii")) is None:
            r_fonts.set(qn("w:ascii"), "Cambria Math")
        if r_fonts.get(qn("w:hAnsi")) is None:
            r_fonts.set(qn("w:hAnsi"), "Cambria Math")
        if not color:
            continue
        w_color = w_rpr.find(_wtag("color"))
        if w_color is None:
            w_color = etree.Element(qn("w:color"))
            w_rpr.append(w_color)
        w_color.set(qn("w:val"), color)


def _apply_native_nary_control_properties(omml_root, *, color: str | None) -> None:
    """
    Integral/summation glyphs are controlled by m:naryPr/m:ctrlPr.
    Add Word-native ctrlPr so operator color follows scoped math color.
    """
    root = omml_root.getroot() if hasattr(omml_root, "getroot") else omml_root
    for nary_pr in root.xpath(".//m:naryPr", namespaces=_OMML_NS):
        ctrl_pr = nary_pr.find(f"{{{OMML_NAMESPACE}}}ctrlPr")
        if ctrl_pr is None:
            ctrl_pr = etree.Element(f"{{{OMML_NAMESPACE}}}ctrlPr")
            nary_pr.append(ctrl_pr)
        w_rpr = ctrl_pr.find(_wtag("rPr"))
        if w_rpr is None:
            w_rpr = etree.Element(qn("w:rPr"))
            ctrl_pr.append(w_rpr)
        r_fonts = w_rpr.find(_wtag("rFonts"))
        if r_fonts is None:
            r_fonts = etree.Element(qn("w:rFonts"))
            w_rpr.append(r_fonts)
        if r_fonts.get(qn("w:ascii")) is None:
            r_fonts.set(qn("w:ascii"), "Cambria Math")
        if r_fonts.get(qn("w:hAnsi")) is None:
            r_fonts.set(qn("w:hAnsi"), "Cambria Math")
        if not color:
            continue
        w_color = w_rpr.find(_wtag("color"))
        if w_color is None:
            w_color = etree.Element(qn("w:color"))
            w_rpr.append(w_color)
        w_color.set(qn("w:val"), color)


def _wtag(local: str) -> str:
    return f"{{{WORD_NAMESPACE}}}{local}"


def _coerce_omml_node(omml_root):
    if hasattr(omml_root, "getroot"):
        return omml_root.getroot()
    return omml_root


def _wrap_display_omath(node):
    # Word-native block equations are serialized as m:oMathPara.
    if node.tag == f"{{{OMML_NAMESPACE}}}oMathPara":
        return node
    if node.tag != f"{{{OMML_NAMESPACE}}}oMath":
        return node
    para = etree.Element(f"{{{OMML_NAMESPACE}}}oMathPara")
    para.append(node)
    return para
