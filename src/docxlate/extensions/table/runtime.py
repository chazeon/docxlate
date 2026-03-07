from __future__ import annotations

import re

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Mm, Pt
from plasTeX import Command, Environment

from docxlate.model import RenderContext
from docxlate.registry import MacroSpec


class table(Environment):
    args = "[ position ]"


class tabular(Environment):
    args = "colspec"


class multicolumn(Command):
    args = "cols:str align:str self"


def _node_name(node) -> str:
    node_name = getattr(node, "nodeName", None)
    return str(node_name).lstrip("\\") if node_name is not None else ""


def _split_tabular_rows(nodes: list) -> list[list[list]]:
    rows: list[list[list]] = [[]]
    current_cell: list = []

    for child in nodes:
        raw_name = getattr(child, "nodeName", None)
        if raw_name is not None and str(raw_name) == "\\":
            rows[-1].append(current_cell)
            current_cell = []
            rows.append([])
            continue
        if _node_name(child) == "active::&":
            rows[-1].append(current_cell)
            current_cell = []
            continue
        current_cell.append(child)

    rows[-1].append(current_cell)

    while rows and all(len(cell) == 0 for cell in rows[-1]):
        rows.pop()
    if not rows:
        return []
    return rows


def _alignment_tokens_from_colspec(colspec: str) -> list[str]:
    return [item["align"] for item in _colspec_descriptors(colspec)]


def _alignment_for_token(token: str):
    if token == "c":
        return WD_ALIGN_PARAGRAPH.CENTER
    if token == "r":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _read_braced(text: str, start: int) -> tuple[str, int]:
    i = start
    while i < len(text) and text[i].isspace():
        i += 1
    if i >= len(text) or text[i] != "{":
        return "", i
    depth = 0
    i += 1
    begin = i
    while i < len(text):
        ch = text[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            if depth == 0:
                return text[begin:i], i + 1
            depth -= 1
        i += 1
    return text[begin:], len(text)


def _parse_length_emu(value: str, *, textwidth_emu: int | None) -> int | None:
    cleaned = str(value or "").strip()
    if not cleaned:
        return None

    m = re.fullmatch(
        r"([0-9]*\.?[0-9]+)\s*(\\textwidth|in|cm|mm|pt)",
        cleaned,
    )
    if m is None:
        return None
    scalar = float(m.group(1))
    unit = m.group(2)
    if unit == "\\textwidth":
        if textwidth_emu is None:
            return None
        return int(scalar * textwidth_emu)
    if unit == "in":
        return int(Inches(scalar))
    if unit == "cm":
        return int(Cm(scalar))
    if unit == "mm":
        return int(Mm(scalar))
    if unit == "pt":
        return int(Pt(scalar))
    return None


def _colspec_descriptors(colspec: str) -> list[dict[str, str | int | None]]:
    out: list[dict[str, str | int | None]] = []
    text = str(colspec or "")
    i = 0
    while i < len(text):
        ch = text[i]
        if ch.isspace() or ch == "|":
            i += 1
            continue
        if ch in {"l", "c", "r"}:
            out.append({"align": ch, "width_spec": None})
            i += 1
            continue
        if ch in {"p", "m", "b"}:
            width_spec, nxt = _read_braced(text, i + 1)
            out.append({"align": "l", "width_spec": width_spec.strip() if width_spec else None})
            i = max(i + 1, nxt)
            continue
        i += 1
    return out


def _resolve_table_style(doc, *, candidates: list[str], fallback: str | None) -> str | None:
    table_styles = [s for s in doc.styles if s.type == WD_STYLE_TYPE.TABLE]
    if not table_styles:
        return None

    style_by_name = {s.name: s for s in table_styles}
    style_by_id = {str(s.style_id): s for s in table_styles}

    for candidate in candidates:
        style = style_by_name.get(candidate) or style_by_id.get(candidate)
        if style is not None:
            return str(style.name)

    if fallback:
        style = style_by_name.get(fallback) or style_by_id.get(fallback)
        if style is not None:
            return str(style.name)

    # deterministic final fallback: first table style in template
    return str(table_styles[0].name)


def _render_caption(latex, caption_node, *, label_names: list[str], anchor_paragraph):
    refs = latex.context.get("refs", {})
    resolved = None
    for label_name in label_names:
        info = refs.get(label_name, {})
        value = info.get("ref_num")
        if value is not None:
            resolved = str(value)
            break

    caption_ctx = RenderContext().with_para_role("caption")
    caption_para = latex.add_paragraph_for_role("caption")
    prefix = f"Table {resolved or '?'}"
    with latex.render_frame(paragraph=caption_para, style=caption_ctx):
        latex.append_inline(f"{prefix}. ", style={"bold": True, "theme": "major"})
        fragment = getattr(caption_node, "attributes", {}).get("self")
        if fragment is not None and getattr(fragment, "childNodes", None):
            latex.render_nodes(fragment.childNodes)
        else:
            text = latex.get_arg_text(caption_node, 0, key="self")
            if text:
                latex.append_inline(text)

    resolver = getattr(latex, "reference_resolver", None)
    if resolver is None:
        return
    current = latex._current_paragraph
    latex._current_paragraph = caption_para
    try:
        for label_name in label_names:
            resolver.register_label(
                latex,
                label_name,
                ref_text=resolved,
            )
    finally:
        latex._current_paragraph = current


def register(latex, *, plugin):
    def _coerce_multicolumn(cell_nodes: list) -> tuple[list, int, str | None]:
        meaningful = []
        for node in cell_nodes:
            if isinstance(node, str):
                if node.strip():
                    meaningful.append(node)
                continue
            if _node_name(node) == "#text" and not str(node).strip():
                continue
            meaningful.append(node)

        if len(meaningful) != 1:
            return cell_nodes, 1, None

        node = meaningful[0]
        if _node_name(node) != "multicolumn":
            return cell_nodes, 1, None

        span_text = latex.get_arg_text(node, 0, key="cols")
        align = latex.get_arg_text(node, 1, key="align")
        try:
            span = int(span_text)
        except (TypeError, ValueError):
            span = 1
        if span < 1:
            span = 1
        fragment = getattr(node, "attributes", {}).get("self")
        if fragment is not None and getattr(fragment, "childNodes", None):
            return list(fragment.childNodes), span, align
        return list(getattr(node, "childNodes", []) or []), span, align

    def _render_tabular(node):
        rows = _split_tabular_rows(list(getattr(node, "childNodes", []) or []))
        if not rows:
            return None

        colspec = latex.get_arg_text(node, 0, key="colspec")
        descriptors = _colspec_descriptors(colspec)
        align_tokens = _alignment_tokens_from_colspec(colspec)
        row_models: list[list[tuple[list, int, str | None]]] = []
        max_cols = 0
        for row in rows:
            row_model = []
            row_total = 0
            for cell_nodes in row:
                content_nodes, span, align = _coerce_multicolumn(cell_nodes)
                row_model.append((content_nodes, span, align))
                row_total += span
            row_models.append(row_model)
            max_cols = max(max_cols, row_total)

        col_count = max(max_cols, len(align_tokens))
        if col_count <= 0:
            return None

        latex._flush_paragraph()
        doc_table = latex.doc.add_table(rows=len(row_models), cols=col_count)
        selected_style = _resolve_table_style(
            latex.doc,
            candidates=plugin.style_candidates(latex),
            fallback=plugin.fallback_style(latex),
        )
        if selected_style is not None:
            try:
                doc_table.style = selected_style
            except Exception:
                latex.context.setdefault("warnings", []).append(
                    f"Failed to apply table style: {selected_style}"
                )
        doc_table.autofit = plugin.autofit(latex)

        textwidth_emu = None
        try:
            section = latex.doc.sections[-1]
            textwidth_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
        except Exception:
            textwidth_emu = None

        has_explicit_width = False
        for idx, desc in enumerate(descriptors):
            if idx >= col_count:
                break
            width_spec = desc.get("width_spec")
            if not isinstance(width_spec, str) or not width_spec:
                continue
            width_emu = _parse_length_emu(width_spec, textwidth_emu=textwidth_emu)
            if width_emu is None or width_emu <= 0:
                continue
            has_explicit_width = True
            doc_table.columns[idx].width = width_emu
            for row_idx in range(len(row_models)):
                doc_table.cell(row_idx, idx).width = width_emu
        if has_explicit_width:
            doc_table.autofit = False

        anchor_paragraph = None
        for row_idx, row_cells in enumerate(row_models):
            col_idx = 0
            for content_nodes, span, cell_align in row_cells:
                if col_idx >= col_count:
                    break
                start_cell = doc_table.cell(row_idx, col_idx)
                end_col = min(col_count - 1, col_idx + span - 1)
                cell = (
                    start_cell.merge(doc_table.cell(row_idx, end_col))
                    if end_col > col_idx
                    else start_cell
                )
                paragraph = cell.paragraphs[0]
                anchor_paragraph = paragraph
                if cell_align:
                    align_token = next(
                        (ch for ch in cell_align.lower() if ch in {"c", "l", "r"}),
                        None,
                    )
                    if align_token is not None:
                        paragraph.alignment = _alignment_for_token(align_token)
                elif col_idx < len(align_tokens):
                    paragraph.alignment = _alignment_for_token(align_tokens[col_idx])
                with latex.render_frame(paragraph=paragraph):
                    latex.render_nodes(content_nodes)
                col_idx += max(1, span)
        latex._flush_paragraph()
        return anchor_paragraph

    @latex.env("tabular", parse_class=tabular)
    def handle_tabular(node):
        _render_tabular(node)

    @latex.env("table", parse_class=table)
    def handle_table(node):
        caption_node = None
        label_names: list[str] = []
        tabular_nodes: list = []
        passthrough_nodes: list = []

        for child in list(getattr(node, "childNodes", []) or []):
            name = _node_name(child)
            if name == "caption":
                caption_node = child
                continue
            if name == "label":
                label_name = latex.get_arg_text(child, 0, key="label").strip()
                if label_name:
                    label_names.append(label_name)
                continue
            if name == "tabular":
                tabular_nodes.append(child)
                continue
            passthrough_nodes.append(child)

        anchor_paragraph = None
        for tabular_node in tabular_nodes:
            rendered_anchor = _render_tabular(tabular_node)
            if rendered_anchor is not None:
                anchor_paragraph = rendered_anchor

        if caption_node is not None:
            _render_caption(
                latex,
                caption_node,
                label_names=label_names,
                anchor_paragraph=anchor_paragraph,
            )
        elif label_names:
            resolver = getattr(latex, "reference_resolver", None)
            if resolver is not None and anchor_paragraph is not None:
                current = latex._current_paragraph
                latex._current_paragraph = anchor_paragraph
                try:
                    for label_name in label_names:
                        resolver.register_label(latex, label_name)
                finally:
                    latex._current_paragraph = current

        if passthrough_nodes:
            latex.render_nodes(passthrough_nodes)

    latex.register_spec(
        MacroSpec(
            name="multicolumn",
            kind="command",
            parse_class=multicolumn,
            policy="stub",
        )
    )
    return None


__all__ = ["register"]
