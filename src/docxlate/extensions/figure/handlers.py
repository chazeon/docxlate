from __future__ import annotations

from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from docxlate.model import RenderContext

from .captioning import CAPTION_SLOT_TOKEN, render_caption_with_template
from .geometry.image import resolve_image_path, resolve_target_width_emu
from .geometry.wrap import wrapped_figure_box_size
from .layout.anchor_host import trim_trailing_whitespace_runs
from .layout.caption_box import estimate_caption_box_height_emu
from .macros import caption, docxlatefigwrapset, includegraphics, wrapfigure


def _current_wrap_entry(latex) -> dict | None:
    stack = latex.context.get("figure_stack", [])
    if stack and stack[-1].get("kind") == "wrapfigure":
        return stack[-1]
    return None


def _parse_directive_inches(value: str) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _apply_wrap_directive(entry: dict, path: str, inches: float) -> bool:
    if path == "figure.wrap.shift.y":
        entry["pos_y_emu"] = int(entry.get("pos_y_emu", 0)) + int(Inches(inches))
        return True
    if path == "figure.wrap.gap":
        if inches < 0:
            return False
        entry["gap_emu"] = int(Inches(inches))
        return True

    side_to_dist = {
        "left": "dist_l",
        "right": "dist_r",
        "top": "dist_t",
        "bottom": "dist_b",
    }
    side_to_inset = {
        "left": "l_ins",
        "right": "r_ins",
        "top": "t_ins",
        "bottom": "b_ins",
    }
    if path.startswith("figure.wrap.pad."):
        side = path.rsplit(".", 1)[-1]
        if side not in side_to_dist or inches < 0:
            return False
        distances = entry.setdefault("wrap_distances_emu", {})
        distances[side_to_dist[side]] = int(Inches(inches))
        return True
    if path.startswith("figure.wrap.inset."):
        side = path.rsplit(".", 1)[-1]
        if side not in side_to_inset or inches < 0:
            return False
        insets = entry.setdefault("textbox_insets_emu", {})
        insets[side_to_inset[side]] = int(Inches(inches))
        return True
    return False


def register_handlers(latex, *, plugin):
    @latex.command(
        "docxlatefigwrapset",
        inline=True,
        parse_class=docxlatefigwrapset,
    )
    def handle_docxlatefigwrapset(node):
        path = latex.get_arg_text(node, 0, key="path").strip().lower()
        value = latex.get_arg_text(node, 1, key="value")
        inches = _parse_directive_inches(value)
        if inches is None:
            latex.context.setdefault("warnings", []).append(
                f"Invalid docxlate figure wrap directive value for {path}: {value!r}"
            )
            return
        entry = _current_wrap_entry(latex)
        if entry is None:
            latex.context.setdefault("warnings", []).append(
                f"Ignored docxlate directive outside wrapfigure: {path}"
            )
            return
        if not _apply_wrap_directive(entry, path, inches):
            latex.context.setdefault("warnings", []).append(
                f"Invalid or unsupported docxlate wrap directive: {path}={value}"
            )

    @latex.command("includegraphics", inline=True, parse_class=includegraphics)
    def handle_includegraphics(node):
        raw_path = latex.get_arg_text(node, 0, key="file")
        image_path = resolve_image_path(latex, raw_path)
        if image_path is None:
            latex.context.setdefault("warnings", []).append(
                f"Image file not found: {raw_path}"
            )
            latex.append_inline(f"[Missing image: {raw_path}]")
            return

        paragraph = latex._active_paragraph()
        if paragraph is None:
            latex._ensure_paragraph()
            paragraph = latex._active_paragraph()
        if paragraph is None:
            return

        stack = latex.context.get("figure_stack", [])
        in_wrapfigure = bool(stack and stack[-1].get("kind") == "wrapfigure")
        target_width_emu = resolve_target_width_emu(latex, node, stack if in_wrapfigure else [])

        try:
            run = latex.emit_image(paragraph, str(image_path), width_emu=target_width_emu)
        except Exception:
            run = latex.emit_image(paragraph, str(image_path))

        if in_wrapfigure:
            drawing = run._r.find(qn("w:drawing"))
            inline = drawing.find(qn("wp:inline")) if drawing is not None else None
            extent = inline.find(qn("wp:extent")) if inline is not None else None
            if extent is not None and stack:
                rendered_cx = int(extent.get("cx", "0"))
                rendered_cy = int(extent.get("cy", "0"))
                stack[-1]["image_run"] = run
                stack[-1]["image_cx_emu"] = rendered_cx
                stack[-1]["image_cy_emu"] = rendered_cy
                # Keep group width synced with actual rendered picture width.
                stack[-1]["target_cx_emu"] = rendered_cx if rendered_cx > 0 else target_width_emu

    @latex.command("caption", inline=True, parse_class=caption)
    def handle_caption(node):
        stack = latex.context.get("figure_stack", [])
        caption_ctx = RenderContext().with_para_role("caption")
        p = latex.add_paragraph_for_role("caption")
        templated = render_caption_with_template(latex, node, plugin=plugin)
        if templated is not None:
            self_fragment = getattr(node, "attributes", {}).get("self")
            prefix, sep, suffix = templated.partition(CAPTION_SLOT_TOKEN)
            if prefix:
                latex.render_latex_fragment(prefix, paragraph=p, style=caption_ctx)
            with latex.render_frame(paragraph=p, style=caption_ctx):
                if self_fragment is not None and getattr(self_fragment, "childNodes", None):
                    latex.render_nodes(self_fragment.childNodes)
                else:
                    text = latex.get_arg_text(node, 0, key="self")
                    latex.append_inline(text)
            if sep and suffix:
                latex.render_latex_fragment(suffix, paragraph=p, style=caption_ctx)
        else:
            self_fragment = getattr(node, "attributes", {}).get("self")
            with latex.render_frame(paragraph=p, style=caption_ctx):
                if self_fragment is not None and getattr(self_fragment, "childNodes", None):
                    latex.render_nodes(self_fragment.childNodes)
                else:
                    text = latex.get_arg_text(node, 0, key="self")
                    latex.append_inline(text)

        if stack and stack[-1].get("kind") == "wrapfigure":
            image_run = stack[-1].get("image_run")
            gap_emu = int(stack[-1].get("gap_emu", plugin.caption_gap_emu(latex)))
            anchor_mode = stack[-1].get("caption_anchor_mode", plugin.wrap_caption_anchor_mode(latex))
            pos_y_emu = int(stack[-1].get("pos_y_emu", plugin.wrap_offset_y_emu(latex)))
            wrap_distances_emu = stack[-1].get("wrap_distances_emu")
            textbox_insets_emu = stack[-1].get("textbox_insets_emu")
            box_cx, _ = wrapped_figure_box_size(stack[-1], caption_cy=0, gap_emu=gap_emu)
            caption_cy = estimate_caption_box_height_emu(p.text, box_cx)
            box_cx, box_cy = wrapped_figure_box_size(stack[-1], caption_cy=caption_cy, gap_emu=gap_emu)
            image_cy = int(stack[-1].get("image_cy_emu", 1000000))

            if image_run is not None and anchor_mode != "separate":
                latex.emitter.emit_wrapped_figure_caption_group_anchor(
                    latex.doc,
                    image_run=image_run,
                    caption_paragraph=p,
                    anchor_paragraph=stack[-1].get("anchor_paragraph"),
                    place=stack[-1].get("place"),
                    pos_y_emu=pos_y_emu,
                    box_cx_emu=box_cx,
                    box_cy_emu=box_cy,
                    gap_emu=gap_emu,
                    wrap_distances_emu=wrap_distances_emu,
                    textbox_insets_emu=textbox_insets_emu,
                )
                stack[-1]["wrapped_emitted"] = True
            else:
                latex.emitter.emit_wrapped_caption_anchor(
                    latex.doc,
                    source_paragraph=p,
                    anchor_paragraph=stack[-1].get("anchor_paragraph"),
                    place=stack[-1].get("place"),
                    pos_y_emu=pos_y_emu + image_cy + gap_emu,
                    box_cx_emu=box_cx,
                    box_cy_emu=caption_cy,
                    wrap_distances_emu=wrap_distances_emu,
                    textbox_insets_emu=textbox_insets_emu,
                )
                if image_run is None:
                    stack[-1]["wrapped_emitted"] = True

    @latex.env("wrapfigure", parse_class=wrapfigure)
    def handle_wrapfigure(node):
        lines = latex.get_arg_text(node, 0, key="lines")
        place = latex.get_arg_text(node, 0, key="place")
        width = latex.get_arg_text(node, 1, key="width")
        created_anchor_host = False
        if latex.doc.paragraphs:
            p = latex.doc.paragraphs[-1]
        else:
            p = latex.add_paragraph_for_role("body")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            created_anchor_host = True
        stack = latex.context.setdefault("figure_stack", [])
        base_pos_y_emu = plugin.wrap_offset_y_emu(latex)
        pos_y_emu = int(base_pos_y_emu)
        stack.append(
            {
                "kind": "wrapfigure",
                "place": place,
                "width": width,
                "lines": lines,
                "pos_y_emu": pos_y_emu,
                "caption_anchor_mode": plugin.wrap_caption_anchor_mode(latex),
                "anchor_paragraph": p,
            }
        )
        try:
            previous_suppress = bool(latex.context.get("_suppress_whitespace_text", False))
            latex.context["_suppress_whitespace_text"] = True
            try:
                with latex.render_frame(paragraph=p):
                    latex.render_nodes(node.childNodes)
            finally:
                latex.context["_suppress_whitespace_text"] = previous_suppress
            if stack and stack[-1].get("kind") == "wrapfigure":
                image_run = stack[-1].get("image_run")
                if image_run is not None and not stack[-1].get("wrapped_emitted"):
                    anchor = latex.emitter.convert_image_run_to_wrap_anchor(
                        image_run,
                        place=stack[-1].get("place"),
                        pos_y_emu=int(stack[-1].get("pos_y_emu", plugin.wrap_offset_y_emu(latex))),
                        wrap_distances_emu=stack[-1].get("wrap_distances_emu"),
                    )
                    if anchor is not None:
                        stack[-1]["wrapped_emitted"] = True
            trim_trailing_whitespace_runs(p)
            if created_anchor_host:
                # If wrapfigure appears at the top before any body text exists,
                # reuse this host for the immediate following paragraph content.
                latex._current_paragraph = p
                latex.context["_skip_next_par_break_once"] = True
                latex.context["_preserve_paragraph_after_env_once"] = True
        finally:
            stack.pop()

    return None


__all__ = ["register_handlers"]
