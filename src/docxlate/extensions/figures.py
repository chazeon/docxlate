from __future__ import annotations

from pathlib import Path

from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from jinja2 import Environment as JinjaEnvironment
from plasTeX import Command, Environment

from docxlate.model import RenderContext


class includegraphics(Command):
    args = "[ options ] file:str"


class caption(Command):
    args = "[ toc ] self"


class wrapfigure(Environment):
    args = "[ lines ] place:str width"


def _resolve_image_path(latex, raw_path: str) -> Path | None:
    if not raw_path:
        return None
    candidate = Path(raw_path)
    if candidate.is_absolute() and candidate.exists():
        return candidate

    bases = []
    tex_path = latex.context.get("tex_path")
    if tex_path:
        bases.append(Path(tex_path).resolve().parent)
    bases.append(Path.cwd())

    for base in bases:
        p = (base / candidate).resolve()
        if p.exists():
            return p
        for ext in (".png", ".jpg", ".jpeg", ".pdf"):
            with_ext = p.with_suffix(ext)
            if with_ext.exists():
                return with_ext
    return None


def _section_textwidth_emu(latex) -> int:
    try:
        section = latex.doc.sections[-1]
        width = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
        return max(width, int(Inches(4.0)))
    except Exception:
        return int(Inches(6.0))


def _parse_latex_length_inches(raw_value: str | None) -> float | None:
    if not raw_value:
        return None
    value = raw_value.strip()
    if value.endswith("in"):
        try:
            return float(value[:-2])
        except ValueError:
            return None
    if value.endswith("cm"):
        try:
            return float(value[:-2]) / 2.54
        except ValueError:
            return None
    if value.endswith("pt"):
        try:
            return float(value[:-2]) / 72.0
        except ValueError:
            return None
    return None


def _parse_textwidth_fraction_emu(raw_value: str | None, textwidth_emu: int) -> int | None:
    if not raw_value:
        return None
    value = raw_value.strip().replace(" ", "")
    if not value.endswith("\\textwidth"):
        return None
    try:
        return int(float(value[: -len("\\textwidth")]) * textwidth_emu)
    except ValueError:
        return None


def _parse_latex_length_emu(raw_value: str | None, textwidth_emu: int) -> int | None:
    if not raw_value:
        return None
    by_textwidth = _parse_textwidth_fraction_emu(raw_value, textwidth_emu)
    if by_textwidth is not None:
        return by_textwidth
    inches = _parse_latex_length_inches(raw_value)
    if inches is not None:
        return int(Inches(inches))
    return None


def _resolve_width_hint(latex, node):
    options = latex.get_arg_text(node, 0, key="options")
    width_hint = None
    if options and "width=" in options:
        for token in options.split(","):
            token = token.strip()
            if token.startswith("width="):
                width_hint = token.split("=", 1)[1].strip()
                break
    return width_hint


def _resolve_target_width_emu(latex, node, stack) -> int:
    textwidth_emu = _section_textwidth_emu(latex)
    width_hint = _resolve_width_hint(latex, node)
    if width_hint:
        parsed = _parse_latex_length_emu(width_hint, textwidth_emu)
        if parsed and parsed > 0:
            return parsed
    if stack:
        wrap_width = stack[-1].get("width")
        parsed = _parse_latex_length_emu(wrap_width, textwidth_emu)
        if parsed and parsed > 0:
            return parsed
    return int(Inches(4.5))


def _estimate_caption_box_height_emu(text: str, box_cx_emu: int) -> int:
    caption = (text or "").strip()
    if not caption:
        return 320000
    # Approximate line capacity from box width using ~5.2 pt average glyph width.
    width_pt = max(40.0, box_cx_emu / 12700.0)
    chars_per_line = max(16, int(width_pt / 5.2))
    line_count = max(1, (len(caption) + chars_per_line - 1) // chars_per_line)
    line_height_emu = int(12 * 12700)  # ~12pt
    padding = int(8 * 12700)
    return max(320000, line_count * line_height_emu + padding)


def _trim_trailing_whitespace_runs(paragraph):
    # Remove trailing whitespace-only text runs on anchor host paragraphs.
    while paragraph.runs:
        run = paragraph.runs[-1]
        text = run.text or ""
        if not text or not text.isspace():
            break
        run._r.getparent().remove(run._r)


def _caption_gap_emu(latex) -> int:
    value = latex.context.get("wrapfigure_caption_gap_in")
    if value is None:
        return 114300
    try:
        inches = float(value)
    except (TypeError, ValueError):
        return 114300
    if inches < 0:
        return 114300
    return int(Inches(inches))


def _caption_template_env() -> JinjaEnvironment:
    return JinjaEnvironment(
        autoescape=False,
        trim_blocks=True,
        lstrip_blocks=True,
        variable_start_string="<<",
        variable_end_string=">>",
        block_start_string="<%",
        block_end_string="%>",
    )


def _caption_template_source(template: str) -> str:
    source = template
    # Allow simple {{ var }} placeholders as a compatibility shorthand.
    if "{{" in source and "<<" not in source:
        source = source.replace("{{", "<<").replace("}}", ">>")
    return source


def _caption_tex_from_node(latex, node) -> str:
    fragment = getattr(node, "attributes", {}).get("self")
    source = getattr(fragment, "source", None) if fragment is not None else None
    if source is not None and str(source).strip():
        return str(source).strip()
    return latex.get_arg_text(node, 0, key="self")


def _fragment_text(value) -> str | None:
    if value is None:
        return None
    text_content = getattr(value, "textContent", None)
    if text_content is not None:
        text = str(text_content).strip()
        if text:
            return text
    source = getattr(value, "source", None)
    if source is not None:
        text = str(source).strip()
        if text:
            return text
    text = str(value).strip()
    if text and "plasTeX.TeXFragment object" not in text:
        return text
    return None


def _find_figure_label(latex, node) -> str | None:
    scope = getattr(node, "parentNode", None)
    while scope is not None:
        stack = list(getattr(scope, "childNodes", []) or [])
        while stack:
            child = stack.pop(0)
            if getattr(child, "nodeName", None) == "label":
                label_name = latex.get_arg_text(child, 0, key="label")
                if label_name:
                    return label_name
            stack[0:0] = list(getattr(child, "childNodes", []) or [])
        node_name = getattr(scope, "nodeName", None)
        if node_name in {"figure", "wrapfigure"}:
            break
        scope = getattr(scope, "parentNode", None)
    return None


def _resolved_label_number(latex, label_name: str | None) -> str:
    if not label_name:
        return "?"
    refs = latex.context.get("refs", {})
    ref_info = refs.get(label_name, {})
    ref_text = ref_info.get("ref_num")
    if ref_text is not None:
        return str(ref_text)
    labels = latex.context.get("labels", {})
    known = labels.get(label_name, {})
    known_text = known.get("ref_text")
    if known_text:
        return str(known_text)
    return "?"


def _figure_name_from_node(node) -> str:
    name = _fragment_text(getattr(node, "captionName", None))
    if name and "\\" not in name and "{" not in name and "}" not in name:
        return name
    return "Figure"


def _figure_number_from_node(node) -> str | None:
    value = _fragment_text(getattr(node, "ref", None))
    if not value:
        return None
    if "\\" in value or "{" in value or "}" in value:
        return None
    return value


def _render_caption_with_template(latex, node) -> str | None:
    template = latex.context.get("figure_caption_template")
    if not template:
        return None
    slot = "__DOCXLATE_CAPTION_SLOT__"
    label_name = _find_figure_label(latex, node)
    fig_num = _resolved_label_number(latex, label_name)
    if fig_num == "?":
        parsed_ref = _figure_number_from_node(node)
        if parsed_ref:
            fig_num = parsed_ref
    fig_name = _figure_name_from_node(node)
    env = _caption_template_env()
    compiled = env.from_string(_caption_template_source(str(template)))
    return compiled.render(
        x=fig_num,
        number=fig_num,
        fig_num=fig_num,
        thefigure=fig_num,
        fig_name=fig_name,
        figurename=fig_name,
        caption=slot,
        caption_tex=slot,
        label=label_name or "",
    ).strip()


def register(latex):
    for macro_name, macro_class in {
        "includegraphics": includegraphics,
        "caption": caption,
        "wrapfigure": wrapfigure,
    }.items():
        latex.macro(macro_name, macro_class)

    @latex.command("includegraphics", inline=True)
    def handle_includegraphics(node):
        raw_path = latex.get_arg_text(node, 0, key="file")
        image_path = _resolve_image_path(latex, raw_path)
        if image_path is None:
            latex.context.setdefault("warnings", []).append(
                f"Image file not found: {raw_path}"
            )
            latex.append_inline(f"[Missing image: {raw_path}]")
            return

        paragraph = latex._active_paragraph()
        if paragraph is None:
            latex._ensure_paragraph()
            paragraph = latex._active_paragraph()
        if paragraph is None:
            return

        stack = latex.context.get("figure_stack", [])
        in_wrapfigure = bool(stack and stack[-1].get("kind") == "wrapfigure")
        target_width_emu = _resolve_target_width_emu(latex, node, stack if in_wrapfigure else [])

        try:
            run = latex.emit_image(paragraph, str(image_path), width_emu=target_width_emu)
        except Exception:
            run = latex.emit_image(paragraph, str(image_path))

        if in_wrapfigure:
            drawing = run._r.find(qn("w:drawing"))
            inline = drawing.find(qn("wp:inline")) if drawing is not None else None
            extent = inline.find(qn("wp:extent")) if inline is not None else None
            if extent is not None and stack:
                stack[-1]["image_run"] = run
                stack[-1]["image_cx_emu"] = int(extent.get("cx", "0"))
                stack[-1]["image_cy_emu"] = int(extent.get("cy", "0"))
                stack[-1]["target_cx_emu"] = target_width_emu

    @latex.command("caption", inline=True)
    def handle_caption(node):
        stack = latex.context.get("figure_stack", [])
        caption_ctx = RenderContext().with_para_role("caption")
        p = latex.add_paragraph_for_role("caption")
        templated = _render_caption_with_template(latex, node)
        if templated is not None:
            slot = "__DOCXLATE_CAPTION_SLOT__"
            self_fragment = getattr(node, "attributes", {}).get("self")
            prefix, sep, suffix = templated.partition(slot)
            if prefix:
                latex.render_latex_fragment(prefix, paragraph=p, style=caption_ctx)
            with latex.render_frame(paragraph=p, style=caption_ctx):
                if self_fragment is not None and getattr(self_fragment, "childNodes", None):
                    latex.render_nodes(self_fragment.childNodes)
                else:
                    text = latex.get_arg_text(node, 0, key="self")
                    latex.append_inline(text)
            if sep and suffix:
                latex.render_latex_fragment(suffix, paragraph=p, style=caption_ctx)
        else:
            self_fragment = getattr(node, "attributes", {}).get("self")
            with latex.render_frame(paragraph=p, style=caption_ctx):
                if self_fragment is not None and getattr(self_fragment, "childNodes", None):
                    latex.render_nodes(self_fragment.childNodes)
                else:
                    text = latex.get_arg_text(node, 0, key="self")
                    latex.append_inline(text)
        if stack and stack[-1].get("kind") == "wrapfigure":
            image_cx = int(stack[-1].get("image_cx_emu", 2160000))
            image_cy = int(stack[-1].get("image_cy_emu", 1000000))
            box_cx = int(stack[-1].get("target_cx_emu", image_cx))
            image_run = stack[-1].get("image_run")
            caption_cy = _estimate_caption_box_height_emu(p.text, box_cx)
            gap_emu = _caption_gap_emu(latex)
            if image_run is not None:
                latex.emit_wrapped_figure_caption_group_anchor(
                    image_run=image_run,
                    caption_paragraph=p,
                    anchor_paragraph=stack[-1].get("anchor_paragraph"),
                    place=stack[-1].get("place"),
                    pos_y_emu=0,
                    box_cx_emu=max(1200000, box_cx),
                    box_cy_emu=max(320000, image_cy + gap_emu + caption_cy),
                    gap_emu=gap_emu,
                )
                stack[-1]["wrapped_emitted"] = True
            else:
                latex.emit_wrapped_caption_anchor(
                    source_paragraph=p,
                    anchor_paragraph=stack[-1].get("anchor_paragraph"),
                    place=stack[-1].get("place"),
                    pos_y_emu=image_cy + gap_emu,
                    box_cx_emu=max(1200000, box_cx),
                    box_cy_emu=caption_cy,
                )
                stack[-1]["wrapped_emitted"] = True

    @latex.env("wrapfigure")
    def handle_wrapfigure(node):
        lines = latex.get_arg_text(node, 0, key="lines")
        place = latex.get_arg_text(node, 0, key="place")
        width = latex.get_arg_text(node, 1, key="width")
        created_anchor_host = False
        if latex.doc.paragraphs:
            p = latex.doc.paragraphs[-1]
        else:
            p = latex.add_paragraph_for_role("body")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            created_anchor_host = True
        stack = latex.context.setdefault("figure_stack", [])
        stack.append(
            {
                "kind": "wrapfigure",
                "place": place,
                "width": width,
                "lines": lines,
                "anchor_paragraph": p,
            }
        )
        try:
            previous_suppress = bool(latex.context.get("_suppress_whitespace_text", False))
            latex.context["_suppress_whitespace_text"] = True
            try:
                with latex.render_frame(paragraph=p):
                    latex.render_nodes(node.childNodes)
            finally:
                latex.context["_suppress_whitespace_text"] = previous_suppress
            if stack and stack[-1].get("kind") == "wrapfigure":
                image_run = stack[-1].get("image_run")
                if image_run is not None and not stack[-1].get("wrapped_emitted"):
                    anchor = latex.convert_image_run_to_wrap_anchor(
                        image_run,
                        place=stack[-1].get("place"),
                        pos_y_emu=0,
                    )
                    if anchor is not None:
                        stack[-1]["wrapped_emitted"] = True
            _trim_trailing_whitespace_runs(p)
            if created_anchor_host:
                # If wrapfigure appears at the top before any body text exists,
                # reuse this host for the immediate following paragraph content.
                latex._current_paragraph = p
                latex.context["_skip_next_par_break_once"] = True
                latex.context["_preserve_paragraph_after_env_once"] = True
        finally:
            stack.pop()

    return None
