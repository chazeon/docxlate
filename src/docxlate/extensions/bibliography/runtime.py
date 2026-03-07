from __future__ import annotations

from pathlib import Path

from docx.shared import Inches, RGBColor
from plasTeX import Command

from .artifacts.aux import parse_aux_artifacts
from .artifacts.bbl import format_bibliography_entry, parse_bbl
from .artifacts.bcf import parse_bcf


class DocxlateBibEntry(Command):
    macroName = "docxlatebibentry"
    args = "idx:str self"


class cite(Command):
    args = "bibkeys:str"


BIBLIOGRAPHY_MACRO_DEFAULTS: dict[str, str] = {
    "bibrangedash": "\u2013",
    "bibinitperiod": ".",
    "bibnamedelima": " ",
    "bibinitdelim": " ",
    "bibinithyphendelim": "-",
    "textendash": "\u2013",
    "textemdash": "\u2014",
}


def _reference_number_for_key(key: str, index: int, cite_order: dict[str, int]) -> int:
    value = cite_order.get(key)
    if value is not None:
        return int(value)
    return index + 1


def _parse_positive_int(value: str) -> int | None:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return None
    if parsed <= 0:
        return None
    return parsed


def _compress_numeric_cite_items(
    items: list[tuple[str, str]],
    *,
    min_run: int,
) -> list[tuple[str | None, str]]:
    numeric: list[tuple[str, int]] = []
    for label, value in items:
        num = _parse_positive_int(value)
        if num is None:
            return [(lbl, val) for lbl, val in items]
        numeric.append((label, num))

    numeric.sort(key=lambda pair: pair[1])
    deduped: list[tuple[str, int]] = []
    seen_nums: set[int] = set()
    for label, num in numeric:
        if num in seen_nums:
            continue
        seen_nums.add(num)
        deduped.append((label, num))

    out: list[tuple[str | None, str]] = []
    i = 0
    while i < len(deduped):
        start_label, start_num = deduped[i]
        j = i
        while j + 1 < len(deduped) and deduped[j + 1][1] == deduped[j][1] + 1:
            j += 1
        run_len = j - i + 1
        if run_len >= min_run:
            out.append((start_label, f"{start_num}\u2013{deduped[j][1]}"))
        else:
            for k in range(i, j + 1):
                out.append((deduped[k][0], str(deduped[k][1])))
        i = j + 1
    return out


def register(latex, *, plugin):
    for macro_name in BIBLIOGRAPHY_MACRO_DEFAULTS:
        macro_class = type(
            f"BibMacro_{macro_name}",
            (Command,),
            {"macroName": macro_name, "args": ""},
        )

        @latex.command(macro_name, inline=True, parse_class=macro_class)
        def _handle_bibliography_macro(_node, _macro_name=macro_name):
            latex.append_inline(plugin.macro_text(latex, _macro_name))

    @latex.command("cite", inline=True, parse_class=cite)
    def handle_cite(node):
        cite_order = latex.context.get("cite_order", {})
        refs = latex.context.get("refs", {})
        bib_links = latex.context.get("bib_entry_labels", {})
        bbl_entries = latex.context.get("bbl_entries", {})
        resolver = getattr(latex, "reference_resolver", None)

        label_text = latex.get_arg_text(node, 0, key="bibkeys")
        labels = [lbl.strip() for lbl in label_text.split(",") if lbl.strip()]
        resolved: list[tuple[str, str]] = []
        for label in labels:
            ref_num = cite_order.get(label)
            if ref_num is None:
                ref_info = refs.get(label, {})
                ref_num = ref_info.get("ref_num")
            resolved.append((label, str(ref_num) if ref_num is not None else label))

        seen_values: set[str] = set()
        unique_resolved: list[tuple[str, str]] = []
        for label, value in resolved:
            if value in seen_values:
                continue
            seen_values.add(value)
            unique_resolved.append((label, value))

        citation = plugin.citation_settings(latex)
        compress_ranges = bool(citation["compress_ranges"])
        min_run = int(citation["min_run"])
        cite_items: list[tuple[str | None, str]]
        if compress_ranges:
            cite_items = _compress_numeric_cite_items(unique_resolved, min_run=min_run)
        else:
            cite_items = [(label, value) for label, value in unique_resolved]

        latex.append_inline("[")
        for idx, (label, value) in enumerate(cite_items):
            if idx:
                latex.append_inline(",")
            target_label = None
            if label is not None:
                target_label = bib_links.get(label)
                if target_label is None and label in bbl_entries:
                    target_label = f"bib:{label}"
            if resolver is not None and target_label:
                with latex.render_frame(link={"anchor": resolver.anchor_name(target_label)}):
                    latex.append_inline(value)
            else:
                latex.append_inline(value)
        latex.append_inline("]")

    @latex.command(
        "docxlatebibentry",
        inline=True,
        parse_class=DocxlateBibEntry,
    )
    def handle_docxlate_bib_entry(node):
        idx_text = latex.get_arg_text(node, 0, key="idx")
        try:
            idx = int(idx_text)
        except (TypeError, ValueError):
            return

        ordered_keys = latex.context.get("_bib_render_order", [])
        if idx < 0 or idx >= len(ordered_keys):
            return
        key = ordered_keys[idx]
        cite_order = latex.context.get("_bib_render_cite_order", {})
        layout = latex.context.get("_bib_render_layout") or plugin.layout_settings(latex)
        missing_keys = set(latex.context.get("_bib_missing_keys", set()))
        missing_policy = str(latex.context.get("_bib_missing_policy", "key"))
        is_missing = key in missing_keys

        p = latex.add_paragraph_for_role("bibliography")
        if layout["numbering"] == "bracket":
            indent = Inches(float(layout["indent_in"]))
            p.paragraph_format.left_indent = indent
            p.paragraph_format.first_line_indent = -indent
            p.paragraph_format.tab_stops.add_tab_stop(indent)
            ref_num = _reference_number_for_key(key, idx, cite_order)
            number_run = p.add_run(f"[{ref_num}]\t")
            if is_missing:
                number_run.bold = True
                number_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        if is_missing and missing_policy == "key":
            missing_run = p.add_run(key)
            missing_run.bold = True
            missing_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        else:
            text_fragment = getattr(node, "attributes", {}).get("self")
            if text_fragment is not None and getattr(text_fragment, "childNodes", None):
                with latex.render_frame(paragraph=p):
                    latex.render_nodes(text_fragment.childNodes)

        resolver = getattr(latex, "reference_resolver", None)
        if resolver is not None:
            label_name = f"bib:{key}"
            current = latex._current_paragraph
            latex._current_paragraph = p
            resolver.register_label(latex, label_name)
            latex._current_paragraph = current
            latex.context.setdefault("bib_entry_labels", {})[key] = label_name

    @latex.on("load")
    def on_load(_tex_source, _soup):
        tex_path = latex.context.get("tex_path")
        if not tex_path:
            return
        latex.context["cite_order"] = {}
        latex.context["bbl_entries"] = {}
        latex.context["bib_entry_labels"] = {}
        aux_path = Path(tex_path).with_suffix(".aux")
        if aux_path.exists():
            aux_cache = latex.context.setdefault("_aux_artifacts_cache", {})
            cache_key = str(aux_path.resolve())
            cached = aux_cache.get(cache_key)
            if cached is None:
                refs, _bibcites, cite_order = parse_aux_artifacts(aux_path)
                cached = {"refs": refs, "cite_order": cite_order}
                aux_cache[cache_key] = cached
            latex.context["cite_order"] = dict(cached.get("cite_order", {}))
        bcf_path = Path(tex_path).with_suffix(".bcf")
        if bcf_path.exists() and not latex.context.get("cite_order"):
            latex.context["cite_order"] = parse_bcf(bcf_path)
        bbl_path = Path(tex_path).with_suffix(".bbl")
        if bbl_path.exists():
            latex.context["bbl_entries"] = parse_bbl(bbl_path)

    @latex.on("post_process")
    def append_references():
        entries = latex.context.get("bbl_entries", {})
        cite_order = latex.context.get("cite_order", {})
        missing_policy = plugin.missing_entry_policy(latex)

        if not entries and not cite_order:
            return

        if cite_order:
            ordered_by_cite = [k for k, _ in sorted(cite_order.items(), key=lambda item: item[1])]
            ordered_keys = ordered_by_cite
        else:
            ordered_keys = sorted(entries.keys())
        if not ordered_keys:
            return

        heading = latex.add_paragraph_for_role("references_heading")
        with latex.render_frame(paragraph=heading):
            latex.append_inline("References", style={"bold": True, "theme": "major"})
        latex.mark_next_body_paragraph_first()
        latex.context["bib_entry_labels"] = {}
        layout = plugin.layout_settings(latex)
        bib_template = plugin.template(latex)
        et_al_limit = plugin.et_al_limit(latex)

        chunks: list[str] = []
        missing_keys: set[str] = set()
        for index, key in enumerate(ordered_keys):
            if key in entries:
                text = format_bibliography_entry(
                    entries[key],
                    template=bib_template,
                    et_al_limit=et_al_limit,
                )
            else:
                latex.context.setdefault("warnings", []).append(
                    f"Missing bibliography entry in .bbl: {key}"
                )
                missing_keys.add(key)
                text = key if missing_policy == "key" else ""
            chunks.append(rf"\docxlatebibentry{{{index}}}{{{text}}}")

        latex.context["_bib_render_order"] = ordered_keys
        latex.context["_bib_render_layout"] = layout
        latex.context["_bib_render_cite_order"] = cite_order
        latex.context["_bib_missing_keys"] = missing_keys
        latex.context["_bib_missing_policy"] = missing_policy
        try:
            latex.render_latex_fragment("\n".join(chunks))
        finally:
            latex.context.pop("_bib_render_order", None)
            latex.context.pop("_bib_render_layout", None)
            latex.context.pop("_bib_render_cite_order", None)
            latex.context.pop("_bib_missing_keys", None)
            latex.context.pop("_bib_missing_policy", None)

    return None


__all__ = ["BIBLIOGRAPHY_MACRO_DEFAULTS", "register"]
