from __future__ import annotations

from collections import defaultdict
from contextlib import contextmanager
from io import IOBase
import os
import re
from typing import Callable, Mapping

from docx import Document
from docx.shared import Pt
from plasTeX import Command, Environment
from plasTeX.DOM import Text
from plasTeX.TeX import TeX
from plasTeX.Tokenizer import Tokenizer

from .docx_ext import DocxEmitterBackend
from .model import EquationSpec, LinkTarget, RenderContext, SpanCompositor
from .registry import MacroSpec, normalize_macro_name, validate_macro_spec


class DocxlateDirectiveTokenizer(Tokenizer):
    _directive_re = re.compile(
        r"^\s*docxlate:\s*(?P<path>[A-Za-z0-9_.-]+)\s*=\s*(?P<value>[-+]?\d*\.?\d+)\s*$",
        flags=re.I,
    )

    def __init__(self, source, context, *, directive_rules=None):
        super().__init__(source, context)
        self._directive_rules = list(directive_rules or [])
        self._source_readline = self.readline
        self.readline = self._readline_with_directives

    def _readline_with_directives(self):
        line = self._source_readline()
        if not line:
            return line
        match = self._directive_re.match(line.strip())
        if match is not None:
            path = match.group("path").lower()
            value = match.group("value")
            for path_re, macro_name in self._directive_rules:
                if path_re.fullmatch(path):
                    injected = rf"\{macro_name}{{{path}}}{{{value}}}"
                    self._charBuffer[:0] = list(injected)
                    break
        return line


class DocxlateTeX(TeX):
    def __init__(self, *args, directive_rules=None, **kwargs):
        super().__init__(*args, **kwargs)
        self._directive_rules = list(directive_rules or [])

    def input(self, source):
        if source is None:
            return
        if self.jobname is None:
            if isinstance(source, str):
                self.jobname = ""
            elif isinstance(source, IOBase):
                self.jobname = os.path.basename(os.path.splitext(source.name)[0])
        tokenizer = DocxlateDirectiveTokenizer(
            source,
            self.ownerDocument.context,
            directive_rules=self._directive_rules,
        )
        self.inputs.append((tokenizer, iter(tokenizer)))
        self.currentInput = self.inputs[-1]
        return self


class LatexBridge:
    def __init__(self, template_path=None, *, strict_macro_specs: bool = True):
        self.doc = Document(template_path) if template_path else Document()
        self.strict_macro_specs = bool(strict_macro_specs)
        self.command_handlers = {}
        self.env_handlers = {}
        self.event_handlers: dict[str, list] = defaultdict(list)
        self.macro_handlers: dict[str, type] = {}
        self.macro_specs: dict[str, MacroSpec] = {}
        self._directive_rules: list[tuple[re.Pattern[str], str]] = []
        self._parse_skip_initial_policies: list[Callable] = []
        self._parse_skip_retry_policies: list[Callable] = []
        self.context: dict = {}  # For storing state across handlers/events
        self._current_paragraph = None
        self._render_stack: list[dict] = []
        self.style_table = {
            "title": ["Title", "Heading 1", "Normal"],
            "author": ["Author", "Subtitle", "Normal"],
            "date": ["Date", "Subtitle", "Normal"],
            "first_body": [
                "FirstParagraph",
                "First Paragraph",
                "BodyText",
                "Body Text",
                "Normal",
            ],
            "body": ["BodyText", "Body Text", "Normal"],
            "references_heading": [
                "ReferencesHeading",
                "References Heading",
                "Bibliography Heading",
                "Heading1",
                "Heading 1",
            ],
            "bibliography": [
                "Bibliography",
                "References",
                "Reference",
                "BodyText",
                "Body Text",
                "Normal",
            ],
            "equation": ["Equation", "Normal"],
            "caption": ["Caption", "Normal"],
        }
        self._next_body_role = "first_body"
        self._inline_styles = {
            "textbf": {"bold": True},
            "textit": {"italic": True},
            "textup": {"italic": False},
            "textmd": {"bold": False},
            "textsl": {"italic": True},
            "textnormal": {
                "bold": False,
                "italic": False,
                "small_caps": False,
                "monospace": False,
            },
            "emph": {"italic": True},
            "textsc": {"small_caps": True},
            "texttt": {"monospace": True},
        }
        self._declaration_styles = {
            "bfseries": {"bold": True},
            "itshape": {"italic": True},
            "slshape": {"italic": True},
            "scshape": {"small_caps": True},
            "ttfamily": {"monospace": True},
            "upshape": {"italic": False},
            "mdseries": {"bold": False},
            "normalfont": {
                "bold": False,
                "italic": False,
                "small_caps": False,
                "monospace": False,
            },
        }
        self.emitter = DocxEmitterBackend(self.context)
        self._span_compositor = SpanCompositor()

    def reset_document(self, template_path=None, *, keep_template_content=False):
        self.doc = Document(template_path) if template_path else Document()
        if template_path and not keep_template_content:
            self._clear_main_content()
        self._current_paragraph = None
        self._next_body_role = "first_body"
        self.emitter = DocxEmitterBackend(self.context)
        self._span_compositor = SpanCompositor()

    def _clear_main_content(self):
        body = self.doc._element.body
        for child in list(body):
            # Keep section properties from the template; remove existing content.
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)

    def command(
        self,
        name,
        *,
        inline=False,
        parse_class=None,
        policy="render",
    ):
        """Decorator: @app.command('section')"""

        def decorator(f):
            if parse_class is None and policy == "render":
                if self.strict_macro_specs:
                    raise ValueError(
                        f"Decorator registration for command {name!r} requires parse_class "
                        "under strict MacroSpec mode"
                    )
                self.command_handlers[name] = (f, inline)
                return f
            self.register_spec(
                MacroSpec(
                    name=name,
                    kind="command",
                    parse_class=parse_class,
                    handler=f,
                    inline=inline,
                    policy=policy,
                )
            )
            return f

        return decorator

    def env(self, name, *, parse_class=None, policy="render"):
        """Decorator: @app.env('equation')"""

        def decorator(f):
            if parse_class is None and policy == "render":
                if self.strict_macro_specs:
                    raise ValueError(
                        f"Decorator registration for environment {name!r} requires parse_class "
                        "under strict MacroSpec mode"
                    )
                self.env_handlers[name] = f
                return f
            self.register_spec(
                MacroSpec(
                    name=name,
                    kind="env",
                    parse_class=parse_class,
                    handler=f,
                    policy=policy,
                )
            )
            return f

        return decorator

    def on(self, event_name):
        """Decorator: @app.on('post_process')"""

        def decorator(f):
            self.event_handlers[event_name].append(f)
            return f

        return decorator

    def macro(self, name, macro_class):
        """Register a plasTeX macro class for parse-time argument handling."""
        self.macro_handlers[name] = macro_class

    def register_spec(self, spec: MacroSpec):
        validate_macro_spec(spec)
        name = normalize_macro_name(spec.name)
        normalized = MacroSpec(
            name=name,
            kind=spec.kind,
            parse_class=spec.parse_class,
            handler=spec.handler,
            inline=spec.inline,
            policy=spec.policy,
        )

        existing_spec = self.macro_specs.get(name)
        if existing_spec is not None:
            raise ValueError(
                f"Duplicate MacroSpec registration for {name!r} "
                f"(existing kind={existing_spec.kind!r}, new kind={normalized.kind!r})"
            )

        if normalized.kind == "command" and name in self.env_handlers:
            raise ValueError(
                f"Conflicting registration for {name!r}: already registered as environment handler"
            )
        if normalized.kind == "env" and name in self.command_handlers:
            raise ValueError(
                f"Conflicting registration for {name!r}: already registered as command handler"
            )

        if normalized.parse_class is not None:
            existing_macro = self.macro_handlers.get(name)
            if existing_macro is not None and existing_macro is not normalized.parse_class:
                raise ValueError(
                    f"Conflicting parse-class registration for {name!r}: "
                    f"{existing_macro!r} vs {normalized.parse_class!r}"
                )
            self.macro(name, normalized.parse_class)

        if normalized.policy == "render":
            if normalized.kind == "command":
                self.command_handlers[name] = (normalized.handler, normalized.inline)
            else:
                self.env_handlers[name] = normalized.handler

        self.macro_specs[name] = normalized

    def register_specs(self, specs):
        for spec in specs:
            self.register_spec(spec)

    def register_comment_directive(self, *, path_pattern: str, macro_name: str) -> None:
        normalized_macro = normalize_macro_name(macro_name)
        if not normalized_macro:
            raise ValueError("Directive registration requires a non-empty macro_name")
        try:
            pattern = re.compile(path_pattern, flags=re.I)
        except re.error as exc:
            raise ValueError(f"Invalid directive path_pattern {path_pattern!r}") from exc
        signature = (pattern.pattern, pattern.flags, normalized_macro)
        for existing_pattern, existing_macro in self._directive_rules:
            existing_signature = (
                existing_pattern.pattern,
                existing_pattern.flags,
                existing_macro,
            )
            if signature == existing_signature:
                return
        self._directive_rules.append((pattern, normalized_macro))

    def register_parse_skip_policy(self, *, initial=None, retry=None) -> None:
        if initial is not None:
            self._parse_skip_initial_policies.append(initial)
        if retry is not None:
            self._parse_skip_retry_policies.append(retry)

    def _collect_parse_skip_packages(
        self,
        hooks: list[Callable],
        *,
        tex_source: str,
        configured_skip_packages: set[str],
        parse_error: Exception | None = None,
    ) -> set[str]:
        requested: set[str] = set()
        for hook in hooks:
            candidate = hook(tex_source, configured_skip_packages, parse_error)
            if not candidate:
                continue
            for package_name in candidate:
                normalized = str(package_name).strip()
                if normalized:
                    requested.add(normalized)
        requested.difference_update(configured_skip_packages)
        return requested

    def validate_macro_registry(self):
        for name, spec in self.macro_specs.items():
            validate_macro_spec(spec)
            if spec.policy == "render":
                if spec.parse_class is None:
                    raise ValueError(
                        f"MacroSpec({name}) policy='render' is missing parse_class"
                    )
                if name not in self.macro_handlers:
                    raise ValueError(
                        f"MacroSpec({name}) policy='render' is not wired to parser globals"
                    )
                if spec.kind == "command" and name not in self.command_handlers:
                    raise ValueError(
                        f"MacroSpec({name}) policy='render' is missing command handler wiring"
                    )
                if spec.kind == "env" and name not in self.env_handlers:
                    raise ValueError(
                        f"MacroSpec({name}) policy='render' is missing environment handler wiring"
                    )
                continue
            if spec.parse_class is None:
                raise ValueError(
                    f"MacroSpec({name}) policy={spec.policy!r} is missing parse_class"
                )

    def handle_event(self, event_name, *args, **kwargs):
        """Trigger an event handler if it exists."""
        if event_name in self.event_handlers:
            for handler in self.event_handlers[event_name]:
                handler(*args, **kwargs)

    def run(self, tex_source):
        """The 'app.run()' equivalent: Process the document."""

        self.validate_macro_registry()
        parsed_tree = self._parse_source(tex_source)
        self.context["parser_engine_used"] = "plastex"
        self._render_stack.clear()
        self.handle_event("load", tex_source, parsed_tree)
        self._walk(self._root_nodes(parsed_tree))
        self._flush_paragraph()
        self.handle_event("post_process")

    def _parse_source(self, tex_source):
        configured_skip_packages = {
            str(p).strip()
            for p in (self.context.get("parse_skip_packages") or [])
            if str(p).strip()
        }
        initial_skip_packages = self._collect_parse_skip_packages(
            self._parse_skip_initial_policies,
            tex_source=tex_source,
            configured_skip_packages=configured_skip_packages,
        )
        parse_source = self._sanitize_source_for_parse(
            tex_source,
            extra_skip_packages=initial_skip_packages or None,
        )
        parsed = None
        parse_error = None
        try:
            parsed = self._parse_with_registered_macros(parse_source)
        except Exception as exc:
            parse_error = exc

        retry_skip_packages = self._collect_parse_skip_packages(
            self._parse_skip_retry_policies,
            tex_source=tex_source,
            configured_skip_packages=configured_skip_packages,
            parse_error=parse_error,
        )
        retry_skip_packages.difference_update(initial_skip_packages)
        if parsed is None and retry_skip_packages:
            retry_source = self._sanitize_source_for_parse(
                tex_source, extra_skip_packages=retry_skip_packages
            )
            if retry_source != parse_source:
                try:
                    parsed = self._parse_with_registered_macros(retry_source)
                    parse_source = retry_source
                except Exception:
                    pass

        if parsed is None:
            if "\\begin{document}" in parse_source:
                body = self._extract_document_body(parse_source)
                if body is not None:
                    fallback_parsed = self._parse_with_registered_macros(body)
                    self._append_warning_once(
                        f"Full LaTeX parse failed ({type(parse_error).__name__}); used body-only parse fallback."
                    )
                    return fallback_parsed
            raise parse_error

        if self._looks_like_preamble_only(parsed) and "\\begin{document}" in parse_source:
            body = self._extract_document_body(parse_source)
            if body is not None:
                fallback_parsed = self._parse_with_registered_macros(body)
                self._append_warning_once(
                    "Full LaTeX parse produced no document body; used body-only parse fallback."
                )
                return fallback_parsed
        return parsed

    def _parse_with_registered_macros(self, tex_source):
        tex = DocxlateTeX(directive_rules=self._directive_rules)
        for macro_name, macro_class in self.macro_handlers.items():
            tex.ownerDocument.context.addGlobal(macro_name, macro_class)
        tex.input(tex_source)
        parsed = tex.parse()
        macro_ctx = tex.ownerDocument.context.contexts[-1]
        self.context["_parse_macro_context"] = dict(macro_ctx)
        return parsed

    def _append_warning_once(self, message: str):
        warnings = self.context.setdefault("warnings", [])
        if message not in warnings:
            warnings.append(message)

    def _unknown_macro_policy(self) -> str:
        raw = str(self.context.get("unknown_macro_policy", "")).strip().lower()
        if raw in {"warn", "strict"}:
            return raw
        mode = str(self.context.get("mode", "")).strip().lower()
        if mode == "strict":
            return "strict"
        return "warn"

    def _unknown_macro_allowlist(self) -> set[str]:
        configured = self.context.get("unknown_macro_allowlist") or []
        allowlist: set[str] = set()
        for item in configured:
            name = normalize_macro_name(str(item))
            if name:
                allowlist.add(name)
        return allowlist

    def _handle_unknown_macro(self, *, name: str, kind: str) -> None:
        normalized = normalize_macro_name(name)
        if not normalized:
            return
        if normalized in self.macro_specs:
            return
        if normalized.startswith("active::"):
            return
        if normalized in {"bgroup", "egroup"}:
            return
        if normalized in self._unknown_macro_allowlist():
            return
        message = f"Unknown LaTeX {kind}: \\{normalized}"
        if self._unknown_macro_policy() == "strict":
            raise ValueError(message)
        self._append_warning_once(message)

    def _sanitize_source_for_parse(
        self,
        tex_source: str,
        *,
        extra_skip_packages: set[str] | None = None,
    ) -> str:
        skip_packages = set(
            p.strip()
            for p in (self.context.get("parse_skip_packages") or [])
            if str(p).strip()
        )
        if extra_skip_packages:
            skip_packages.update(
                p.strip() for p in extra_skip_packages if str(p).strip()
            )
        skip_paths = set(
            p.strip()
            for p in (self.context.get("parse_skip_usepackage_paths") or [])
            if str(p).strip()
        )
        if not skip_packages and not skip_paths:
            return tex_source

        usepackage_re = re.compile(
            r"\\usepackage(?P<opts>\s*\[[^\]]*\])?\s*\{(?P<pkgs>[^}]*)\}"
        )
        warnings = self.context.setdefault("warnings", [])

        def _replace(match: re.Match) -> str:
            opts = match.group("opts") or ""
            pkgs_raw = match.group("pkgs")
            packages = [pkg.strip() for pkg in pkgs_raw.split(",") if pkg.strip()]
            if not packages:
                return match.group(0)

            kept: list[str] = []
            removed: list[str] = []
            for pkg in packages:
                if pkg in skip_packages or pkg in skip_paths:
                    removed.append(pkg)
                else:
                    kept.append(pkg)
            if not removed:
                return match.group(0)

            message = (
                f"Skipped usepackage for parser compatibility: {', '.join(removed)}"
            )
            if message not in warnings:
                warnings.append(message)
            if not kept:
                return ""
            return f"\\usepackage{opts}" + "{" + ",".join(kept) + "}"

        return usepackage_re.sub(_replace, tex_source)

    def _looks_like_preamble_only(self, parsed_tree):
        meaningful = False
        for node in self._root_nodes(parsed_tree):
            if self._node_kind(node) == "text":
                if str(node).strip():
                    meaningful = True
                    break
                continue
            name = self._node_name(node)
            if not name or name in {"documentclass", "usepackage", "par", "#text"}:
                continue
            meaningful = True
            break
        return not meaningful

    def _extract_document_body(self, tex_source):
        match = re.search(
            r"\\begin\{document\}(.*?)\\end\{document\}",
            tex_source,
            flags=re.S,
        )
        if not match:
            return None
        return match.group(1)

    def _root_nodes(self, parsed_tree):
        body = getattr(parsed_tree, "document", None) or parsed_tree
        if hasattr(body, "contents"):
            return body.contents
        if hasattr(body, "childNodes"):
            return body.childNodes
        return []

    def _node_name(self, node):
        node_name = getattr(node, "nodeName", None)
        return str(node_name).lstrip("\\") if node_name else ""

    def _node_children(self, node):
        return list(getattr(node, "childNodes", []) or [])

    def _node_kind(self, node):
        if isinstance(node, Text) or isinstance(node, str):
            return "text"
        if isinstance(node, Environment):
            return "environment"
        if isinstance(node, Command):
            return "command"
        return "other"

    def _walk(
        self, nodes, render_context: RenderContext | Mapping[str, object] | None = None
    ) -> RenderContext:
        active_ctx = self._coerce_render_context(render_context)
        for node in nodes:
            kind = self._node_kind(node)

            if kind == "text":
                self._append_node_text(node, active_ctx)
                continue

            raw_name = getattr(node, "nodeName", None)
            if raw_name is not None and str(raw_name) == "\\":
                self._emit_line_break()
                continue

            name = self._node_name(node)
            special_text = self._special_text_for_node(name)
            if special_text is not None:
                self._append_text(special_text, context=active_ctx)
                continue
            if name in {"#document", "document"}:
                active_ctx = self._walk(self._node_children(node), active_ctx)
                continue
            if name == "par":
                if self.context.pop("_skip_next_par_break_once", False):
                    active_ctx = self._walk(self._node_children(node), active_ctx)
                    continue
                paragraph = self._active_paragraph()
                preserve_flag = bool(self.context.get("_preserve_paragraph_once"))
                preserve_current = preserve_flag or (
                    paragraph is not None and len(paragraph.runs) == 0
                )
                if preserve_flag:
                    self.context["_preserve_paragraph_once"] = False
                if not preserve_current:
                    self._flush_paragraph()
                active_ctx = self._walk(self._node_children(node), active_ctx)
                self._flush_paragraph()
                continue

            children = self._node_children(node)
            declaration_style = self._declaration_style_for_node(node, name)
            if declaration_style is not None:
                active_ctx = active_ctx.apply_style_delta(declaration_style)
                if children:
                    self._walk(children, active_ctx)
                continue
            if name in self._inline_styles:
                self._walk(children, active_ctx.apply_style_delta(self._inline_styles[name]))
                continue

            handler_entry = self.command_handlers.get(name)
            if handler_entry:
                handler, inline_mode = handler_entry
                if not inline_mode:
                    self._flush_paragraph()
                with self.render_frame(style=active_ctx):
                    handler(node)
                if not inline_mode:
                    self._walk(children, active_ctx)
                    self._flush_paragraph()
                continue

            env_handler = self.env_handlers.get(name)
            if env_handler:
                self._flush_paragraph()
                with self.render_frame(style=active_ctx):
                    env_handler(node)
                preserve_after = bool(
                    self.context.pop("_preserve_paragraph_after_env_once", False)
                )
                if not preserve_after:
                    self._flush_paragraph()
                continue

            if kind == "command":
                self._handle_unknown_macro(name=name, kind="command")
            elif kind == "environment":
                self._handle_unknown_macro(name=name, kind="environment")
            elif kind == "other":
                self._handle_unknown_macro(name=name, kind="command")
            self._walk(children, active_ctx)
        return active_ctx

    def render_nodes(self, nodes, style=None):
        self._walk(list(nodes or []), render_context=style)

    def render_latex_fragment(self, tex_source: str, *, paragraph=None, style=None):
        """
        Parse and render a LaTeX fragment through the regular node walker.

        Important parser-context pitfall:
        Parsing a raw top-level fragment (tex.input(fragment)) does not always
        behave like normal document content parsing in plasTeX. In particular,
        token handling can diverge from content that appears inside command
        arguments in the main document pipeline (for example, dash/ligature
        behavior observed in templates).

        To keep fragment rendering aligned with regular LaTeX content parsing,
        we wrap the fragment in a temporary macro argument and render that
        parsed argument node tree. This forces the same argument-parse context
        while keeping output free of the wrapper command itself.
        """
        class _DocxlateFragment(Command):
            macroName = "docxlatefragment"
            args = "self"

        tex = DocxlateTeX(directive_rules=self._directive_rules)
        macro_context = self.context.get("_parse_macro_context")
        if isinstance(macro_context, dict):
            tex.ownerDocument.context.importMacros(macro_context)
        for macro_name, macro_class in self.macro_handlers.items():
            tex.ownerDocument.context.addGlobal(macro_name, macro_class)
        tex.ownerDocument.context.addGlobal("docxlatefragment", _DocxlateFragment)
        tex.input(r"\docxlatefragment{" + tex_source + "}")
        parsed = tex.parse()
        root_nodes = self._root_nodes(parsed)
        fragment_node = next(
            (
                n
                for n in root_nodes
                if str(getattr(n, "nodeName", "")).lstrip("\\") == "docxlatefragment"
            ),
            None,
        )
        if fragment_node is not None:
            fragment = getattr(fragment_node, "attributes", {}).get("self")
            nodes = list(getattr(fragment, "childNodes", []) or [])
        else:
            nodes = root_nodes
        with self.render_frame(paragraph=paragraph):
            self._walk(nodes, render_context=style)

    @contextmanager
    def render_frame(self, *, paragraph=None, link=None, style=None):
        frame = {}
        if paragraph is not None:
            frame["paragraph"] = paragraph
        if style is not None:
            frame["render_context"] = self._coerce_render_context(style)
        self._render_stack.append(frame)
        link_target = LinkTarget.from_value(link)
        if link_target is None:
            try:
                yield
            finally:
                self._render_stack.pop()
            return
        with self.emitter.link_scope(link_target):
            try:
                yield
            finally:
                self._render_stack.pop()

    def _active_frame_value(self, key):
        for frame in reversed(self._render_stack):
            if key in frame:
                return frame[key]
        return None

    def _active_render_context(self) -> RenderContext:
        current = self._active_frame_value("render_context")
        if isinstance(current, RenderContext):
            return current
        return RenderContext()

    def get_active_render_context(self) -> RenderContext:
        return self._active_render_context()

    def _coerce_render_context(
        self, value: RenderContext | Mapping[str, object] | None
    ) -> RenderContext:
        if isinstance(value, RenderContext):
            return value
        if isinstance(value, Mapping):
            return RenderContext.from_style_mapping(value, fallback=self._active_render_context())
        return self._active_render_context()

    def _special_text_for_node(self, node_name: str):
        # LaTeX tie (`~`) should render as a non-breaking space.
        if node_name == "active::~":
            return "\u00A0"
        literal_map = {
            "%": "%",
            "_": "_",
            "#": "#",
            "&": "&",
            "{": "{",
            "}": "}",
            "textasciitilde": "~",
            "textbackslash": "\\",
            "textquotedblleft": "“",
            "textquotedblright": "”",
            "textquoteleft": "‘",
            "textquoteright": "’",
        }
        if node_name in literal_map:
            return literal_map[node_name]
        return None

    def _declaration_style_for_node(self, node, name: str):
        static = self._declaration_styles.get(name)
        if static is not None:
            return static
        if name != "color":
            return None
        color_spec = self.get_arg_text(node, 0, key="color")
        if not color_spec:
            return None
        return {"color": color_spec}

    def save(self, filename):
        self.doc.save(filename)

    def get_arg_text(self, node, index, default="", key=None):
        args = getattr(node, "args", None)
        if args and not isinstance(args, str) and len(args) > index:
            value = str(args[index]).strip()
            return value.strip("{}")

        attributes = getattr(node, "attributes", None)
        if attributes:
            if key in attributes and attributes[key] is not None:
                return self._stringify_attr_value(attributes[key])
            values = [
                v
                for k, v in attributes.items()
                if v is not None and k not in {"*modifier*", "toc", "len"}
            ]
            if len(values) > index:
                return self._stringify_attr_value(values[index])
        return default

    def _stringify_attr_value(self, value):
        if isinstance(value, list):
            return ",".join(str(item).strip() for item in value if str(item).strip())
        source = getattr(value, "source", None)
        if source is not None:
            return str(source).strip("{} ")
        text_content = getattr(value, "textContent", None)
        if text_content is not None:
            return str(text_content).strip("{} ")
        child_nodes = getattr(value, "childNodes", None)
        if child_nodes:
            text = "".join(str(node) for node in child_nodes if node is not None)
            if text.strip():
                return text.strip("{} ")
        return str(value).strip("{} ")

    def get_node_name(self, node):
        node_name = getattr(node, "nodeName", None)
        if node_name is None:
            return None
        return str(node_name).lstrip("\\")

    def get_node_text(self, node):
        string_attr = getattr(node, "string", None)
        if string_attr is not None:
            text = str(string_attr).strip()
            if text:
                return text

        contents = getattr(node, "contents", None)
        if contents:
            return "".join(str(child) for child in contents if child is not None)

        children = getattr(node, "childNodes", None)
        if children:
            return "".join(str(child) for child in children if child is not None)

        attributes = getattr(node, "attributes", None)
        if attributes:
            text_parts = []
            for value in attributes.values():
                if value is None:
                    continue
                part = self._stringify_attr_value(value)
                if part:
                    text_parts.append(part)
            if text_parts:
                return " ".join(text_parts)

        return str(node)

    def get_math_source(self, node):
        source = getattr(node, "source", None)
        if source:
            source = str(source).strip()
            if source.startswith("\\begin{equation}") and source.endswith(
                "\\end{equation}"
            ):
                source = source[
                    len("\\begin{equation}") : -len("\\end{equation}")
                ].strip()
                return source
            if source.startswith("\\(") and source.endswith("\\)"):
                return source[2:-2].strip()
            if source.startswith("$") and source.endswith("$") and len(source) >= 2:
                return source[1:-1].strip()
            return source

        return self.get_node_text(node)

    def _append_node_text(self, node, render_context: RenderContext):
        text = str(node).replace("\n", " ")
        if self.context.get("_trim_next_leading_space_once"):
            text = text.lstrip(" ")
            self.context["_trim_next_leading_space_once"] = False
        if text == "":
            return
        if self.context.get("_suppress_whitespace_text") and text.isspace():
            return
        paragraph = self._active_paragraph()
        if text.isspace():
            if paragraph is None:
                return
            if not paragraph.runs:
                return
            if paragraph.runs[-1].text.endswith(" "):
                return
            text = " "
        self._append_text(text, context=render_context)

    def _ensure_paragraph(self):
        if self._active_frame_value("paragraph") is not None:
            return
        if self._current_paragraph is None:
            role = self._active_render_context().para_role or self._next_body_role
            self._current_paragraph = self.emitter.begin_paragraph(
                self.doc,
                role=role,
                style_table=self.style_table,
            )
            self._consume_first_line_intent_once(self._current_paragraph)
            if self._active_render_context().para_role is None:
                self._next_body_role = "body"

    def request_noindent(self):
        paragraph = self._active_paragraph()
        if paragraph is not None and not any(run.text.strip() for run in paragraph.runs):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            self.context["_first_line_intent_once"] = None
            return
        self.context["_first_line_intent_once"] = "noindent"

    def request_indent(self):
        paragraph = self._active_paragraph()
        if paragraph is not None and not any(run.text.strip() for run in paragraph.runs):
            paragraph.paragraph_format.first_line_indent = None
            self.context["_first_line_intent_once"] = None
            return
        self.context["_first_line_intent_once"] = "indent"

    def _consume_first_line_intent_once(self, paragraph):
        intent = self.context.get("_first_line_intent_once")
        if not intent:
            return
        if intent == "noindent":
            paragraph.paragraph_format.first_line_indent = Pt(0)
        elif intent == "indent":
            paragraph.paragraph_format.first_line_indent = None
        self.context["_first_line_intent_once"] = None

    def mark_next_body_paragraph_first(self):
        self._next_body_role = "first_body"

    def add_paragraph_for_role(self, role):
        return self.emitter.begin_paragraph(
            self.doc, role=role, style_table=self.style_table
        )

    def _append_text(
        self,
        text,
        style: Mapping[str, object] | None = None,
        *,
        role: str | None = None,
        context: RenderContext | None = None,
    ):
        self._ensure_paragraph()
        paragraph = self._active_paragraph()
        if paragraph is None:
            return
        base_ctx = context or self._active_render_context()
        span = self._span_compositor.compose(
            text,
            base=base_ctx,
            style_delta=style,
            role=role,
        )
        self.emitter.emit_span(paragraph, span)

    def _active_paragraph(self):
        return self._active_frame_value("paragraph") or self._current_paragraph

    def _flush_paragraph(self):
        self._current_paragraph = None

    def append_inline(self, text, style=None):
        self._append_text(text, style)

    def emit_text(self, text, *, style=None, role: str | None = None):
        self._append_text(text, style, role=role)

    def append_math(self, latex_str):
        self._ensure_paragraph()
        paragraph = self._active_paragraph()
        if paragraph is None:
            return
        active_color = self._active_render_context().style.color
        if self.emitter.has_active_link():
            # DOCX hyperlinks don't safely embed OMML; keep math text visible.
            fallback_style = {"theme": "minor"}
            if active_color:
                fallback_style["color"] = active_color
            self._append_text(latex_str, fallback_style, context=RenderContext())
            return
        self.emitter.emit_equation(
            paragraph,
            EquationSpec(
                latex=latex_str,
                color=active_color,
                display=False,
                style=self._active_render_context().style,
            ),
        )

    def emit_equation(
        self,
        latex_str: str,
        *,
        number: str | None = None,
        paragraph=None,
        para_role: str | None = None,
        color: str | None = None,
    ):
        active_color = color if color is not None else self._active_render_context().style.color
        if paragraph is None:
            role = para_role or self._active_render_context().para_role
            paragraph = self.emitter.begin_paragraph(
                self.doc, role=role, style_table=self.style_table
            )
        self.emitter.emit_equation(
            paragraph,
            EquationSpec(
                latex=latex_str,
                number=number,
                color=active_color,
                display=True,
                style=self._active_render_context().style,
            ),
        )
        return paragraph

    def emit_image(self, paragraph, image_path: str, *, width_emu: int | None = None):
        return self.emitter.emit_image(paragraph, image_path, width_emu=width_emu)

    def convert_image_run_to_wrap_anchor(
        self,
        run,
        *,
        place: str | None,
        pos_y_emu: int = 0,
        wrap_distances_emu: dict[str, int] | None = None,
    ):
        return self.emitter.convert_image_run_to_wrap_anchor(
            run,
            place=place,
            pos_y_emu=pos_y_emu,
            wrap_distances_emu=wrap_distances_emu,
        )

    def emit_wrapped_caption_anchor(
        self,
        *,
        source_paragraph,
        anchor_paragraph=None,
        place: str | None,
        pos_y_emu: int,
        box_cx_emu: int,
        box_cy_emu: int,
        wrap_distances_emu: dict[str, int] | None = None,
        textbox_insets_emu: dict[str, int] | None = None,
    ):
        return self.emitter.emit_wrapped_caption_anchor(
            self.doc,
            source_paragraph=source_paragraph,
            anchor_paragraph=anchor_paragraph,
            place=place,
            pos_y_emu=pos_y_emu,
            box_cx_emu=box_cx_emu,
            box_cy_emu=box_cy_emu,
            wrap_distances_emu=wrap_distances_emu,
            textbox_insets_emu=textbox_insets_emu,
        )

    def emit_wrapped_figure_caption_group_anchor(
        self,
        *,
        image_run,
        caption_paragraph,
        anchor_paragraph=None,
        place: str | None,
        pos_y_emu: int,
        box_cx_emu: int,
        box_cy_emu: int,
        gap_emu: int,
        wrap_distances_emu: dict[str, int] | None = None,
        textbox_insets_emu: dict[str, int] | None = None,
    ):
        return self.emitter.emit_wrapped_figure_caption_group_anchor(
            self.doc,
            image_run=image_run,
            caption_paragraph=caption_paragraph,
            anchor_paragraph=anchor_paragraph,
            place=place,
            pos_y_emu=pos_y_emu,
            box_cx_emu=box_cx_emu,
            box_cy_emu=box_cy_emu,
            gap_emu=gap_emu,
            wrap_distances_emu=wrap_distances_emu,
            textbox_insets_emu=textbox_insets_emu,
        )

    def _emit_line_break(self):
        self._ensure_paragraph()
        paragraph = self._active_paragraph()
        if paragraph is None:
            return
        self.emitter.emit_line_break(paragraph)
