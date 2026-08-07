from .core import LatexBridge
from .utils import apply_theme_font, inject_omml
from .aux import parse_refs
from .extensions import (
    register_bibliography_extension,
    register_figures_extension,
    register_hyperref_extension,
    register_lists_extension,
    register_newtx_extension,
    register_table_extension,
    register_xcolor_extension,
)
from .model import RenderContext
from pathlib import Path
from docx.shared import Pt
import re
from plasTeX import Command, Environment
from plasTeX.Base.LaTeX.Math import equation as plastex_equation
from plasTeX.Base.LaTeX.Math import math as plastex_math
from plasTeX.Base.LaTeX.Sectioning import paragraph as plastex_paragraph
from plasTeX.Base.LaTeX.Sectioning import section as plastex_section
from plasTeX.Base.LaTeX.Sectioning import subsection as plastex_subsection
from plasTeX.Base.LaTeX.Sectioning import subsubsection as plastex_subsubsection

latex = LatexBridge()
register_hyperref_extension(latex)
register_lists_extension(latex)
register_newtx_extension(latex)
register_table_extension(latex)
register_xcolor_extension(latex)
register_figures_extension(latex)
register_bibliography_extension(latex)


class And(Command):
    macroName = "and"
    args = ""


class Title(Command):
    macroName = "title"
    args = "self"


class Author(Command):
    macroName = "author"
    args = "self"


class Needspace(Command):
    macroName = "Needspace"
    args = "len:str"


class Date(Command):
    macroName = "date"
    args = "self"


class Maketitle(Command):
    macroName = "maketitle"
    args = ""


class Affiliation(Command):
    macroName = "affiliation"
    args = "self"


class Email(Command):
    macroName = "email"
    args = "[ options ] self"


class Orcidlink(Command):
    macroName = "orcidlink"
    args = "self"


class Abstract(Environment):
    macroName = "abstract"


class Noindent(Command):
    macroName = "noindent"
    args = ""


class Ensuremath(Command):
    macroName = "ensuremath"
    args = "self"


class Indent(Command):
    macroName = "indent"
    args = ""


def _math_node_text(node):
    return latex.get_math_source(node)


def _front_matter_tex(node) -> str:
    attributes = getattr(node, "attributes", {}) or {}
    fragment = attributes.get("self")
    source = getattr(fragment, "source", None) if fragment is not None else None
    if source is None:
        source = latex.get_arg_text(node, 0, key="self")
    return str(source).strip() if source is not None else ""


def _store_front_matter_tex(node, key: str, *, append: bool = False):
    text = _front_matter_tex(node)
    if text:
        context_key = f"_frontmatter_{key}_tex"
        if append:
            latex.context.setdefault(f"_frontmatter_{key}s_tex", []).append(text)
        else:
            latex.context[context_key] = text


def _extract_preamble(tex_source: str) -> str:
    marker = r"\begin{document}"
    idx = tex_source.find(marker)
    if idx == -1:
        return tex_source
    return tex_source[:idx]


def _read_balanced_braces(text: str, start_idx: int) -> tuple[str, int]:
    if start_idx >= len(text) or text[start_idx] != "{":
        return "", start_idx
    depth = 0
    i = start_idx + 1
    out: list[str] = []
    while i < len(text):
        ch = text[i]
        if ch == "\\" and i + 1 < len(text):
            out.append(ch)
            out.append(text[i + 1])
            i += 2
            continue
        if ch == "{":
            depth += 1
            out.append(ch)
            i += 1
            continue
        if ch == "}":
            if depth == 0:
                return "".join(out), i + 1
            depth -= 1
            out.append(ch)
            i += 1
            continue
        out.append(ch)
        i += 1
    return "".join(out), i


def _extract_last_braced_command_argument(source: str, command_name: str) -> str | None:
    # Keep escaped percent signs, strip true comments.
    cleaned = re.sub(r"(?<!\\)%.*", "", source)
    pattern = re.compile(rf"(?<!\\)\\{re.escape(command_name)}\b")
    last_value: str | None = None
    pos = 0
    while True:
        match = pattern.search(cleaned, pos)
        if not match:
            break
        i = match.end()
        while i < len(cleaned) and cleaned[i].isspace():
            i += 1
        if i < len(cleaned) and cleaned[i] == "{":
            value, nxt = _read_balanced_braces(cleaned, i)
            if value.strip():
                last_value = value.strip()
            else:
                last_value = ""
            pos = max(nxt, match.end() + 1)
            continue
        pos = match.end() + 1
    return last_value


def _populate_front_matter_from_source(tex_source: str):
    preamble = _extract_preamble(tex_source)
    for key in ("title", "author", "date"):
        context_key = f"_frontmatter_{key}_tex"
        if latex.context.get(context_key):
            continue
        value = _extract_last_braced_command_argument(preamble, key)
        if value is not None:
            latex.context[context_key] = value


def _normalized_title_policy() -> str:
    raw = str(latex.context.get("title_render_policy", "auto")).lower()
    if raw in {"explicit", "auto", "always"}:
        return raw
    return "auto"


def _front_matter_has_content() -> bool:
    return any(
        latex.context.get(k)
        for k in (
            "_frontmatter_title_tex",
            "_frontmatter_author_tex",
            "_frontmatter_date_tex",
            "_frontmatter_authors_tex",
            "_frontmatter_affiliations_tex",
            "_frontmatter_abstract_nodes",
        )
    )


def _used_body_only_parse_fallback() -> bool:
    warnings = latex.context.get("warnings", [])
    return any("body-only parse fallback" in w for w in warnings)


def _prepend_paragraphs(paragraphs):
    body = latex.doc._element.body
    nodes = [p._p for p in paragraphs if p is not None]
    if not nodes:
        return
    for node in nodes:
        body.remove(node)
    first_idx = 0
    for i, child in enumerate(list(body)):
        if child.tag.endswith("}sectPr"):
            continue
        first_idx = i
        break
    for offset, node in enumerate(nodes):
        body.insert(first_idx + offset, node)


def _emit_front_matter(*, prepend: bool = False):
    if latex.context.get("_frontmatter_rendered"):
        return
    title_tex = latex.context.get("_frontmatter_title_tex")
    authors_tex = latex.context.get("_frontmatter_authors_tex") or []
    if not authors_tex and latex.context.get("_frontmatter_author_tex"):
        authors_tex = [latex.context["_frontmatter_author_tex"]]
    affiliations_tex = latex.context.get("_frontmatter_affiliations_tex") or []
    emails_tex = latex.context.get("_frontmatter_emails_tex") or []
    abstract_nodes = latex.context.get("_frontmatter_abstract_nodes") or []
    date_tex = latex.context.get("_frontmatter_date_tex")
    if not any([title_tex, authors_tex, affiliations_tex, abstract_nodes, date_tex]):
        return

    created = []
    if title_tex:
        p_title = latex.add_paragraph_for_role("title")
        latex.render_latex_fragment(
            title_tex, paragraph=p_title, style={"theme": "major"}
        )
        created.append(p_title)

    if authors_tex:
        p_author = latex.add_paragraph_for_role("author")
        latex.context["_in_maketitle_author"] = True
        try:
            for index, author_tex in enumerate(authors_tex):
                if index:
                    p_author.add_run(", ")
                latex.render_latex_fragment(
                    author_tex, paragraph=p_author, style={"theme": "major"}
                )
        finally:
            latex.context["_in_maketitle_author"] = False
        created.append(p_author)

    for affiliation_tex in affiliations_tex:
        paragraph = latex.add_paragraph_for_role("author")
        latex.render_latex_fragment(affiliation_tex, paragraph=paragraph)
        created.append(paragraph)

    for email_tex in emails_tex:
        paragraph = latex.add_paragraph_for_role("author")
        latex.render_latex_fragment(email_tex, paragraph=paragraph)
        created.append(paragraph)

    if date_tex:
        p_date = latex.add_paragraph_for_role("date")
        latex.render_latex_fragment(date_tex, paragraph=p_date)
        created.append(p_date)

    if abstract_nodes:
        p_abstract = latex.add_paragraph_for_role("body")
        with latex.render_frame(paragraph=p_abstract):
            latex.append_inline("Abstract. ", style={"bold": True})
            latex.render_nodes(abstract_nodes)
        created.append(p_abstract)

    if prepend and created:
        _prepend_paragraphs(created)

    latex.context["_frontmatter_rendered"] = True
    latex.mark_next_body_paragraph_first()


def _handle_section(node):
    cmd = latex.get_node_name(node)
    levels: dict[str, int] = {
        "section": 1,
        "subsection": 2,
        "subsubsection": 3,
    }
    level = levels.get(cmd, 1)
    h = latex.doc.add_heading("", level=level)
    title_node = getattr(node, "attributes", {}).get("title")
    if title_node is not None and getattr(title_node, "childNodes", None):
        with latex.render_frame(paragraph=h):
            latex.render_nodes(title_node.childNodes, style={"theme": "major"})
    else:
        title = latex.get_arg_text(node, 0, key="title")
        run = h.add_run(title)
        apply_theme_font(run, "major")
    latex.mark_next_body_paragraph_first()


@latex.command("section", parse_class=plastex_section)
def handle_section(node):
    _handle_section(node)


@latex.command("subsection", parse_class=plastex_subsection)
def handle_subsection(node):
    _handle_section(node)


@latex.command("subsubsection", parse_class=plastex_subsubsection)
def handle_subsubsection(node):
    _handle_section(node)


@latex.command("title", inline=True, parse_class=Title)
def handle_title(node):
    _store_front_matter_tex(node, "title")


@latex.command("author", inline=True, parse_class=Author)
def handle_author(node):
    _store_front_matter_tex(node, "author", append=True)


@latex.command("affiliation", inline=True, parse_class=Affiliation)
def handle_affiliation(node):
    _store_front_matter_tex(node, "affiliation", append=True)


@latex.command("email", inline=True, parse_class=Email)
def handle_email(node):
    _store_front_matter_tex(node, "email", append=True)


@latex.command("orcidlink", inline=True, parse_class=Orcidlink)
def handle_orcidlink(node):
    value = _front_matter_tex(node)
    if value:
        latex.append_inline(f" (ORCID: {value})")


@latex.env("abstract", parse_class=Abstract)
def handle_abstract(node):
    nodes = list(getattr(node, "childNodes", []) or [])
    if latex.context.get("_frontmatter_maketitle_seen"):
        paragraph = latex.add_paragraph_for_role("body")
        with latex.render_frame(paragraph=paragraph):
            latex.append_inline("Abstract. ", style={"bold": True})
            latex.render_nodes(nodes)
        return
    latex.context["_frontmatter_abstract_nodes"] = nodes


@latex.command("date", inline=True, parse_class=Date)
def handle_date(node):
    _store_front_matter_tex(node, "date")


@latex.command("and", inline=True, parse_class=And)
def handle_and(_node):
    if latex.context.get("_in_maketitle_author"):
        paragraph = latex._active_paragraph()
        if paragraph is not None and paragraph.runs:
            paragraph.runs[-1].text = paragraph.runs[-1].text.rstrip()
        latex.append_inline(", ")
        latex.context["_trim_next_leading_space_once"] = True
    else:
        latex.append_inline(" and ")


@latex.command("maketitle", parse_class=Maketitle)
def handle_maketitle(_node):
    latex.context["_frontmatter_maketitle_seen"] = True
    _emit_front_matter(prepend=False)


@latex.command("noindent", inline=True, parse_class=Noindent)
def handle_noindent(_node):
    latex.request_noindent()


@latex.command("indent", inline=True, parse_class=Indent)
def handle_indent(_node):
    latex.request_indent()


@latex.command("Needspace", inline=True, parse_class=Needspace)
def handle_needspace(_node):
    # Layout hint for TeX pagination; no-op for DOCX output.
    return


@latex.command("paragraph", inline=True, parse_class=plastex_paragraph)
def handle_paragraph(node):
    """Render LaTeX \\paragraph as a run-in heading."""
    current = latex._active_paragraph()
    if current is not None and any(run.text.strip() for run in current.runs):
        latex._flush_paragraph()
    latex._ensure_paragraph()
    paragraph = latex._active_paragraph()
    if paragraph is not None:
        # Keep run-in headings flush-left regardless of template body indent.
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)

    title_node = getattr(node, "attributes", {}).get("title")
    with latex.render_frame():
        if title_node is not None and getattr(title_node, "childNodes", None):
            latex.render_nodes(
                title_node.childNodes,
                style={"bold": True, "theme": "major"},
            )
        else:
            title = latex.get_arg_text(node, 0, key="title")
            latex.append_inline(title, style={"bold": True, "theme": "major"})
        # Force body back to regular weight even if paragraph style is bold.
        latex.append_inline(" ", style={"bold": False, "italic": False})
        # The paragraph body frequently begins with a whitespace text node
        # due to source line breaks; consume one to avoid double spaces.
        latex.context["_trim_next_leading_space_once"] = True
        latex.context["_preserve_paragraph_once"] = True
        latex.render_nodes(node.childNodes, style={"bold": False, "italic": False})


@latex.env("equation", parse_class=plastex_equation)
def handle_math(node):
    source = latex.get_math_source(node)
    resolver = getattr(latex, "reference_resolver", None)
    refs = latex.context.get("refs", {})
    labels = re.findall(r"\\label\{([^}]+)\}", source)
    resolved_number: str | None = None
    for label_name in labels:
        ref_info = refs.get(label_name, {})
        ref_text = ref_info.get("ref_num")
        if resolved_number is None and ref_text is not None:
            resolved_number = str(ref_text)
    source = re.sub(r"\\label\{[^}]+\}", "", source).strip()
    equation_ctx = latex.get_active_render_context().with_para_role("equation")
    with latex.render_frame(style=equation_ctx):
        p = latex.emit_equation(source, number=resolved_number)
    if resolver is not None:
        current_paragraph = latex._current_paragraph
        latex._current_paragraph = p
        try:
            for label_name in labels:
                ref_info = refs.get(label_name, {})
                ref_text = ref_info.get("ref_num")
                resolver.register_label(
                    latex,
                    label_name,
                    ref_text=str(ref_text) if ref_text is not None else None,
                )
        finally:
            latex._current_paragraph = current_paragraph


@latex.on("load")
def on_load(tex_source, soup):
    """Parse .aux data for citation handling when available."""
    for key in (
        "_frontmatter_authors_tex",
        "_frontmatter_affiliations_tex",
        "_frontmatter_emails_tex",
        "_frontmatter_abstract_nodes",
    ):
        latex.context.pop(key, None)
    latex.context["_frontmatter_rendered"] = False
    latex.context["_frontmatter_maketitle_seen"] = False
    _populate_front_matter_from_source(tex_source)
    tex_path = latex.context.get("tex_path")
    if not tex_path:
        return
    latex.context["refs"] = {}
    aux_path = Path(tex_path).with_suffix(".aux")
    if aux_path.exists():
        aux_cache = latex.context.setdefault("_aux_artifacts_cache", {})
        cache_key = str(aux_path.resolve())
        cached = aux_cache.get(cache_key)
        if cached is None:
            refs, bibcites = parse_refs(aux_path)
            cached = {"refs": refs, "bibcites": bibcites}
            aux_cache[cache_key] = cached
        latex.context["refs"] = dict(cached.get("refs", {}))


@latex.on("post_process")
def render_implicit_front_matter():
    if (
        latex.context.get("_frontmatter_maketitle_seen")
        and not _front_matter_has_content()
        and _used_body_only_parse_fallback()
    ):
        warnings = latex.context.setdefault("warnings", [])
        msg = (
            "maketitle was found, but title/author/date metadata was unavailable "
            "after body-only parse fallback; front matter was not rendered."
        )
        if msg not in warnings:
            warnings.append(msg)

    if latex.context.get("_frontmatter_rendered"):
        return
    if not _front_matter_has_content():
        return
    policy = _normalized_title_policy()
    if policy == "explicit":
        return
    # If maketitle was present, front matter should already be emitted in-place.
    if latex.context.get("_frontmatter_maketitle_seen"):
        return
    # Auto/always: emit once at document start when maketitle is absent.
    _emit_front_matter(prepend=True)


@latex.command("$", inline=True, parse_class=plastex_math)
def handle_inline_dollar_math(node):
    latex.append_math(_math_node_text(node))


@latex.command("math", inline=True, parse_class=plastex_math)
def handle_inline_paren_math(node):
    latex.append_math(_math_node_text(node))


@latex.command("ensuremath", inline=True, parse_class=Ensuremath)
def handle_ensuremath(node):
    source = latex.get_arg_text(node, 0, key="self")
    if source:
        latex.append_math(source)
